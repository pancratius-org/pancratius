from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from audit import poetry_stanzas


def test_stanza_source_ignores_pagination_and_fused_signoff(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("first")
    paragraph.add_run().add_break(WD_BREAK.COLUMN)
    paragraph.add_run("second")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("third")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("09.02.2025 Сергей Панкратиус")
    path = tmp_path / "poem.docx"
    document.save(str(path))

    assert poetry_stanzas.expected_groups(path, "Unrelated title") == [2]
