"""Poetry lineation checks at the compiler/output boundary."""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from pancratius import docx_conversion
from pancratius.content_catalog import split_frontmatter

ROOT = Path(__file__).resolve().parents[2]
POETRY = ROOT / "src" / "content" / "poetry"
POEM_MARKDOWN = tuple(sorted(POETRY.glob("*/ru.md")))
FOOTNOTE_DEFINITION = re.compile(r"^\[\^[0-9A-Za-z._-]+\]:\s")
IMAGE_ONLY = re.compile(r"!\[[^]]*]\([^)]+\)")


def _reading_stanza_sizes(body: str) -> tuple[int, ...]:
    lines = body.strip().splitlines()
    lines = lines[: next(
        (index for index, line in enumerate(lines) if FOOTNOTE_DEFINITION.match(line)),
        len(lines),
    )]
    groups = (
        group for group in re.split(r"\n\s*\n", "\n".join(lines)) if group.strip()
    )
    return tuple(
        len([line for line in group.splitlines() if line.strip()])
        for group in groups
        if IMAGE_ONLY.fullmatch(group.strip()) is None
    )


@pytest.mark.parametrize(
    "markdown",
    POEM_MARKDOWN,
    ids=[path.parent.name for path in POEM_MARKDOWN],
)
def test_committed_poem_stanza_topology_matches_fresh_compiler(
    markdown: Path,
    tmp_path: Path,
) -> None:
    frontmatter, committed_body = split_frontmatter(
        markdown.read_text(encoding="utf-8")
    )
    docx = markdown.with_suffix(".docx")
    assert docx.is_file(), f"missing source DOCX for {markdown.relative_to(ROOT)}"

    converted = docx_conversion.convert_single_docx(
        docx,
        kind="poem",
        lang="ru",
        work_key=markdown.parent.name,
        title=str(frontmatter["title"]),
        title_index={},
        media_out=tmp_path / "media",
    )

    assert _reading_stanza_sizes(converted.body) == _reading_stanza_sizes(
        committed_body
    )


def test_poem_compiler_keeps_authored_lines_separate_from_pagination(
    tmp_path: Path,
) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("first")
    paragraph.add_run().add_break(WD_BREAK.COLUMN)
    paragraph.add_run("second")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("third")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("09.02.2025 Сергей Панкратиус")
    docx = tmp_path / "poem.docx"
    document.save(str(docx))

    converted = docx_conversion.convert_single_docx(
        docx,
        kind="poem",
        lang="ru",
        work_key="fixture",
        title="Unrelated title",
        title_index={},
        media_out=tmp_path / "media",
    )

    assert _reading_stanza_sizes(converted.body) == (2,)
