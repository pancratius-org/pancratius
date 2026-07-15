from __future__ import annotations

import shutil
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest

from pancratius import cli, docx_inspect, docx_source, docx_structure, ir, ooxml
from pancratius.docx_inspect import (
    DocxInspectError,
    InspectOptions,
    ParaRow,
    parse_index_range,
)
from pancratius.docx_structure import (
    BlockClaim,
    BodyParagraph,
    CompilerBlockKind,
    ContextParagraph,
    ContextReason,
    ContextRole,
    FoldConflict,
    FoldDecision,
    FoldDisposition,
    ReviewParagraph,
    ReviewReason,
    SourceBlockHit,
    _fold_result,
    _observation,
    fold_decisions,
    observe_structure,
    source_block_hits,
)
from pancratius.passes.pipeline import POST_FOLD_SEAM, Context


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _empty_source() -> docx_source.DocxSourceDocument:
    return docx_source.DocxSourceDocument(Path("source.docx"), ())


def test_read_rows_separates_line_breaks_from_page_breaks(tmp_path: Path) -> None:
    """A `<w:br/>` line break is authored LINEATION (counted in br_count); a page or column
    break is PAGINATION (excluded from br_count, surfaced on page_break_* instead). Conflating
    them made the vision render open a near-blank page on a chapter break (E1, book 36)."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document()
    line = doc.add_paragraph("first")
    line.add_run().add_break(WD_BREAK.LINE)           # an authored line break — lineation
    line.add_run("second")
    pg = doc.add_paragraph("before-break")
    pg.add_run().add_break(WD_BREAK.PAGE)             # an inline page break — pagination
    col = doc.add_paragraph("before-column")
    col.add_run().add_break(WD_BREAK.COLUMN)          # column pagination is not a line either
    after = doc.add_paragraph("on-next-page")
    after.paragraph_format.page_break_before = True   # a pageBreakBefore — pagination
    path = tmp_path / "breaks.docx"
    doc.save(str(path))

    rows = docx_inspect.read_rows(docx_source.read(path))
    line_row = next(r for r in rows if r.text.startswith("first"))
    assert line_row.br_count == 1 and not line_row.page_break_inline
    pg_row = next(r for r in rows if r.text == "before-break")
    assert pg_row.br_count == 0 and pg_row.page_break_inline
    col_row = next(r for r in rows if r.text == "before-column")
    assert col_row.br_count == 0 and col_row.column_break_inline
    after_row = next(r for r in rows if r.text == "on-next-page")
    assert after_row.page_break_before and after_row.br_count == 0

    paragraphs = docx_source.read(path).paragraphs
    line_source = next(p for p in paragraphs if p.text.startswith("first"))
    assert line_source.content.breaks == (docx_source.BreakKind.LINE,)
    assert line_source.content.line_segments == ("first", "second")
    page_source = next(p for p in paragraphs if p.text == "before-break")
    assert page_source.content.breaks == (docx_source.BreakKind.PAGE,)
    assert page_source.content.line_segments == ("before-break",)


def test_render_distinguishes_pagination_only_from_structural_empty(
    tmp_path: Path,
) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK

    document = Document()
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph()
    document.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)
    page_before = document.add_paragraph()
    page_before.paragraph_format.page_break_before = True
    path = tmp_path / "pagination-signals.docx"
    document.save(str(path))

    rendered = docx_inspect.render(
        docx_inspect.read_rows(docx_source.read(path))
    ).splitlines()
    signals = [line for line in rendered if line[:4].strip().isdigit()]

    assert "pagination pageBr" in signals[0]
    assert "pagination" not in signals[1]
    assert "pagination colBr" in signals[2]
    assert "pagination pageBefore" in signals[3]


def test_paragraph_content_derives_every_view_from_one_atom_sequence() -> None:
    content = docx_source.ParagraphContent((
        docx_source.TextAtom("before"),
        docx_source.BreakKind.PAGE,
        docx_source.TextAtom("after"),
        docx_source.BreakKind.LINE,
        docx_source.TextAtom("next"),
        docx_source.TextAtom("\t"),
    ))

    assert content.reading == "before after next"
    assert content.lineated == "before after\nnext"
    assert content.breaks == (
        docx_source.BreakKind.PAGE,
        docx_source.BreakKind.LINE,
    )
    assert content.line_segments == ("before after", "next")

    authored_space = docx_source.ParagraphContent((
        docx_source.TextAtom("a  b\u00a0c"),
        docx_source.BreakKind.LINE,
        docx_source.TextAtom("d  e"),
    ))
    assert authored_space.reading == "a b c d e"
    assert authored_space.line_segments == ("a b c", "d e")



def test_adjudication_fingerprints_encode_domain_equivalence() -> None:
    def content(*atoms: docx_source.ParagraphAtom) -> docx_source.ParagraphContent:
        return docx_source.ParagraphContent(atoms)

    space = content(docx_source.TextAtom("a b"))
    tab = content(docx_source.TextAtom("a\tb"))
    page = content(
        docx_source.TextAtom("a"),
        docx_source.BreakKind.PAGE,
        docx_source.TextAtom("b"),
    )
    column = content(
        docx_source.TextAtom("a"),
        docx_source.BreakKind.COLUMN,
        docx_source.TextAtom("b"),
    )
    line = content(
        docx_source.TextAtom("a"),
        docx_source.BreakKind.LINE,
        docx_source.TextAtom("b"),
    )

    lineation = docx_source.SourceAdjudicationKind.LINEATION
    scripture = docx_source.SourceAdjudicationKind.SCRIPTURE
    assert {
        item.adjudication_fingerprint(lineation)
        for item in (space, tab, page, column)
    } == {space.adjudication_fingerprint(lineation)}
    assert line.adjudication_fingerprint(lineation) != space.adjudication_fingerprint(lineation)
    assert {
        item.adjudication_fingerprint(scripture)
        for item in (space, tab, page, column, line)
    } == {space.adjudication_fingerprint(scripture)}

    composed = content(docx_source.TextAtom("é"))
    decomposed = content(docx_source.TextAtom("e\u0301"))
    assert composed.adjudication_fingerprint(lineation) == decomposed.adjudication_fingerprint(
        lineation
    )
    assert composed.adjudication_fingerprint(scripture) == decomposed.adjudication_fingerprint(
        scripture
    )

@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        (None, docx_source.BreakKind.LINE),
        ("textWrapping", docx_source.BreakKind.LINE),
        ("page", docx_source.BreakKind.PAGE),
        ("column", docx_source.BreakKind.COLUMN),
    ],
)
def test_break_kind_accepts_exactly_the_ooxml_vocabulary(
    raw: str | None,
    expected: docx_source.BreakKind,
) -> None:
    assert docx_source.BreakKind.from_ooxml(raw) is expected


def test_break_kind_fails_closed_on_future_ooxml_syntax() -> None:
    with pytest.raises(
        docx_source.DocxSourceError,
        match=r"unsupported w:br type 'future-layout'",
    ):
        docx_source.BreakKind.from_ooxml("future-layout")


def test_paragraph_markers_keep_independent_source_facts() -> None:
    markers = docx_source._paragraph_markers(
        numbered=True,
        direct_style="Heading1",
        text="---",
    )

    assert markers == docx_source.ParagraphMarkers(
        numbered=True,
        heading_style=True,
        thematic_marker=True,
    )


def test_document_layout_retains_heterogeneous_section_widths() -> None:
    root = ET.fromstring(
        f'<w:document xmlns:w="{ooxml.W_NS}"><w:body>'
        '<w:sectPr><w:pgSz w:w="10000"/><w:pgMar w:left="1000" w:right="1000"/></w:sectPr>'
        '<w:sectPr><w:pgSz w:w="11000"/><w:pgMar w:left="1000" w:right="1000"/></w:sectPr>'
        '</w:body></w:document>'
    )

    layout = docx_source._document_layout(root, docx_source._style_sheet_xml(None))

    assert layout.column_width == docx_source.HeterogeneousColumnWidths(
        (docx_source.Twips(8000), docx_source.Twips(9000))
    )


def test_document_layout_does_not_promote_partial_section_geometry() -> None:
    root = ET.fromstring(
        f'<w:document xmlns:w="{ooxml.W_NS}"><w:body>'
        '<w:sectPr><w:pgSz w:w="10000"/></w:sectPr>'
        '<w:sectPr><w:pgSz w:w="11000"/><w:pgMar w:left="1000" w:right="1000"/></w:sectPr>'
        '</w:body></w:document>'
    )

    layout = docx_source._document_layout(root, docx_source._style_sheet_xml(None))

    assert layout.column_width == docx_source.PartiallyObservedColumnWidths(
        (docx_source.Twips(9000),)
    )


@pytest.mark.parametrize(
    ("content", "page_break_before", "has_opaque_payload", "expected"),
    [
        (
            docx_source.ParagraphContent((docx_source.TextAtom("readable"),)),
            True,
            True,
            docx_source.ParagraphDisposition.CONTENT,
        ),
        (
            docx_source.ParagraphContent(),
            False,
            False,
            docx_source.ParagraphDisposition.STRUCTURAL_EMPTY,
        ),
        (
            docx_source.ParagraphContent(),
            True,
            False,
            docx_source.ParagraphDisposition.PAGINATION_ONLY,
        ),
        (
            docx_source.ParagraphContent((docx_source.BreakKind.PAGE,)),
            False,
            False,
            docx_source.ParagraphDisposition.PAGINATION_ONLY,
        ),
        (
            docx_source.ParagraphContent((docx_source.BreakKind.LINE,)),
            False,
            False,
            docx_source.ParagraphDisposition.NON_TEXT,
        ),
        (
            docx_source.ParagraphContent(),
            False,
            True,
            docx_source.ParagraphDisposition.NON_TEXT,
        ),
    ],
)
def test_paragraph_disposition_is_derived_from_canonical_facts(
    content: docx_source.ParagraphContent,
    *,
    page_break_before: bool,
    has_opaque_payload: bool,
    expected: docx_source.ParagraphDisposition,
) -> None:
    semantics = docx_source.ParagraphSemantics(
        content=content,
        page_break_before=page_break_before,
        payload=docx_source.ParagraphPayload(
            frozenset({docx_source.ParagraphPayloadKind.OPAQUE})
            if has_opaque_payload
            else frozenset()
        ),
    )

    assert semantics.disposition is expected


def test_source_read_reports_unknown_break_with_document_and_paragraph(
    tmp_path: Path,
) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    paragraph = document.add_paragraph("before")
    element = OxmlElement("w:br")
    element.set(qn("w:type"), "future-layout")
    paragraph.add_run()._r.append(element)
    path = tmp_path / "unknown-break.docx"
    document.save(str(path))

    with pytest.raises(
        docx_source.DocxSourceError,
        match=(
            r"unknown-break\.docx: paragraph 0: "
            r"unsupported w:br type 'future-layout'"
        ),
    ):
        docx_source.read(path)


def test_pagination_only_is_excluded_without_breaking_semantic_adjacency(
    tmp_path: Path,
) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK

    document = Document()
    document.add_paragraph("before")
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph()
    document.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    document.add_paragraph("after")
    path = tmp_path / "dispositions.docx"
    document.save(str(path))

    source = docx_source.read(path)
    assert [paragraph.disposition for paragraph in source.paragraphs] == [
        docx_source.ParagraphDisposition.CONTENT,
        docx_source.ParagraphDisposition.PAGINATION_ONLY,
        docx_source.ParagraphDisposition.STRUCTURAL_EMPTY,
        docx_source.ParagraphDisposition.NON_TEXT,
        docx_source.ParagraphDisposition.CONTENT,
    ]
    assert [
        paragraph.reconciliation_position.value
        for paragraph in source.reconciliation_paragraphs
        if paragraph.reconciliation_position is not None
    ] == [0, 1, 2, 3]
    assert source.paragraphs[1].reconciliation_position is None


pandoc_required = pytest.mark.skipif(
    shutil.which("pandoc") is None,
    reason="pandoc is required for importer-backed DOCX inspection",
)


@pandoc_required
def test_docx_inspect_cli_smoke_with_temp_docx_fixture(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, ["Alpha opening", "Beta marker", "Gamma close"])

    rc = cli.main(["docx", "inspect", str(source)])

    assert rc == 0
    out = capsys.readouterr().out
    assert "body paragraphs" in out
    assert "Alpha opening" in out
    assert "Beta marker" in out


@pandoc_required
def test_docx_inspect_cli_contains_filter(tmp_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, ["Alpha opening", "Beta marker", "Gamma close"])

    rc = cli.main(["docx", "inspect", str(source), "--contains", "Beta"])

    assert rc == 0
    out = capsys.readouterr().out
    assert "3 body paragraphs, 1 shown" in out
    assert "Beta marker" in out
    assert "Alpha opening" not in out
    assert "Gamma close" not in out


def test_docx_inspect_cli_missing_file_is_usage_error(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    missing = tmp_path / "missing.docx"

    rc = cli.main(["docx", "inspect", str(missing)])

    assert rc == 2
    assert "DOCX not found" in capsys.readouterr().err


def test_docx_inspect_cli_reports_source_error_without_traceback(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    source = tmp_path / "source.docx"
    source.write_bytes(b"source boundary stub")

    def fail_read(_path: Path) -> docx_source.DocxSourceDocument:
        raise docx_source.DocxSourceError(
            "source.docx: paragraph 4: unsupported w:br type 'future-layout'"
        )

    monkeypatch.setattr(docx_source, "read", fail_read)

    rc = cli.main(["docx", "inspect", str(source)])

    assert rc == 2
    error = capsys.readouterr().err
    assert "unsupported w:br type 'future-layout'" in error
    assert "Traceback" not in error


def test_docx_inspect_cli_accepts_book_selector(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    resolved = tmp_path / "book-30" / "ru.docx"
    seen: list[Path] = []

    def fake_resolve_book_docx(number: int, *, lang: str, content_root: Path) -> Path:
        assert number == 30
        assert lang == "ru"
        assert content_root == tmp_path
        return resolved

    monkeypatch.setattr(docx_inspect, "resolve_book_docx", fake_resolve_book_docx)

    def fake_inspect(docx: Path, _options: InspectOptions) -> object:
        seen.append(docx)
        return object()

    monkeypatch.setattr(docx_inspect, "inspect_docx", fake_inspect)
    monkeypatch.setattr(docx_inspect, "render_inspection", lambda _result: "resolved book source")

    rc = cli.main(["docx", "inspect", "book:30", "--around", "Alpha", "--content-root", str(tmp_path)])

    assert rc == 0
    assert seen == [resolved]
    assert "resolved book source" in capsys.readouterr().out


def test_docx_inspect_cli_rejects_non_book_selector(capsys: pytest.CaptureFixture[str]) -> None:
    rc = cli.main(["docx", "inspect", "poem:1"])

    assert rc == 2
    assert "not a DOCX source selector" in capsys.readouterr().err


def test_docx_inspect_cli_keeps_unknown_colon_source_as_path(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    seen: list[Path] = []

    def fake_inspect(docx: Path, _options: InspectOptions) -> object:
        seen.append(docx)
        return object()

    monkeypatch.setattr(docx_inspect, "inspect_docx", fake_inspect)
    monkeypatch.setattr(docx_inspect, "render_inspection", lambda _result: "ok")

    rc = cli.main(["docx", "inspect", "notes:v1.docx"])

    assert rc == 0
    assert seen == [Path("notes:v1.docx")]


def test_docx_inspect_rejects_ambiguous_filters() -> None:
    with pytest.raises(DocxInspectError, match="choose only one inspect filter"):
        InspectOptions.from_cli(contains="Alpha", index_range=parse_index_range("0:2"))


def test_docx_inspect_marks_repeated_text_with_mixed_import_roles(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    row = ParaRow(
        index=0,
        text="Same",
        style="Normal",
        direct_style="",
        align="",
        contextual=False,
        spacing={},
        indent={},
        numbered=False,
        border="",
        heading=False,
        thematic=False,
        br_count=0,
        empty=False,
        disposition=docx_source.ParagraphDisposition.CONTENT,
    )

    monkeypatch.setattr(
        docx_inspect,
        "classify_blocks",
        lambda _docx: docx_inspect.BlockClassifications(
            by_text={"Same": frozenset({
                CompilerBlockKind.PARAGRAPH,
                CompilerBlockKind.LINEATED,
            })},
            by_source={},
        ),
    )

    docx_inspect.annotate([row], _empty_source())

    assert row.block_kind == "Ambiguous[LineatedBlock|Paragraph]"


def test_docx_inspect_prefers_source_span_classification(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    row = ParaRow(
        index=4,
        text="Repeated",
        style="Normal",
        direct_style="",
        align="",
        contextual=False,
        spacing={},
        indent={},
        numbered=False,
        border="",
        heading=False,
        thematic=False,
        br_count=0,
        empty=False,
        disposition=docx_source.ParagraphDisposition.CONTENT,
    )
    span = ir.SourceSpan(4, 6)

    monkeypatch.setattr(
        docx_inspect,
        "classify_blocks",
        lambda _docx: docx_inspect.BlockClassifications(
            by_text={"Repeated": frozenset({CompilerBlockKind.PARAGRAPH})},
            by_source={4: docx_inspect.BlockSourceHit((
                BlockClaim(CompilerBlockKind.LINEATED, span, (0,)),
            ))},
        ),
    )

    docx_inspect.annotate([row], _empty_source())

    assert row.block_kind == "LineatedBlock"
    assert row.block_source_span == span


def test_docx_inspect_classifies_empty_rows_inside_source_span(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    row = ParaRow(
        index=5,
        text="",
        style="Normal",
        direct_style="",
        align="",
        contextual=False,
        spacing={},
        indent={},
        numbered=False,
        border="",
        heading=False,
        thematic=False,
        br_count=0,
        empty=True,
        disposition=docx_source.ParagraphDisposition.STRUCTURAL_EMPTY,
    )
    span = ir.SourceSpan(4, 6)

    monkeypatch.setattr(
        docx_inspect,
        "classify_blocks",
        lambda _docx: docx_inspect.BlockClassifications(
            by_text={},
            by_source={5: docx_inspect.BlockSourceHit((
                BlockClaim(CompilerBlockKind.LINEATED, span, (0,)),
            ))},
        ),
    )

    docx_inspect.annotate([row], _empty_source())

    assert row.block_kind == "LineatedBlock"
    assert "ir=4..6" in docx_inspect.render([row])


def test_docx_inspect_does_not_classify_empty_rows_from_enclosing_block_span(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    row = ParaRow(
        index=5,
        text="",
        style="Normal",
        direct_style="",
        align="",
        contextual=False,
        spacing={},
        indent={},
        numbered=False,
        border="",
        heading=False,
        thematic=False,
        br_count=0,
        empty=True,
        disposition=docx_source.ParagraphDisposition.STRUCTURAL_EMPTY,
    )

    def fake_adapt(
        _source: docx_source.DocxSourceDocument,
        _media_dir: Path,
        _diagnostics: list[ir.Diagnostic],
    ) -> ir.Document:
        return ir.Document(blocks=[
            ir.LineatedBlock(
                stanzas=[[ir.Line([ir.Text("before")])], [ir.Line([ir.Text("after")])]],
                register=ir.Register.VERSE,
                source_span=ir.SourceSpan(4, 6),
            )
        ])

    monkeypatch.setattr(docx_inspect.da, "adapt", fake_adapt)

    docx_inspect.annotate([row], _empty_source())

    assert row.block_kind == "—"
    assert row.block_source_span is None


def test_docx_inspect_kind_filters_keep_ambiguous_candidates() -> None:
    rows = [
        ParaRow(
            index=0,
            text="same",
            style="Normal",
            direct_style="",
            align="",
            contextual=False,
            spacing={},
            indent={},
            numbered=False,
            border="",
            heading=False,
            thematic=False,
            br_count=0,
            empty=False,
            disposition=docx_source.ParagraphDisposition.CONTENT,
            block_kind="Ambiguous[LineatedBlock|Paragraph]",
        )
    ]

    selected = docx_inspect.select_rows(rows, InspectOptions.from_cli(lineated_only=True))

    assert selected == rows


# --- total structural observation ----------------------------------------------------------


def _hit(kinds: set[CompilerBlockKind], start: int, end: int) -> SourceBlockHit:
    span = ir.SourceSpan(start, end)
    return SourceBlockHit(tuple(
        BlockClaim(kind=kind, span=span, path=(index,))
        for index, kind in enumerate(kinds)
    ))


def test_structure_observation_is_closed_and_reasoned() -> None:
    assert isinstance(_observation(_hit({CompilerBlockKind.PARAGRAPH}, 5, 5)), BodyParagraph)
    assert isinstance(_observation(_hit({CompilerBlockKind.LINEATED}, 5, 8)), BodyParagraph)
    structural = _observation(_hit({CompilerBlockKind.HEADING}, 3, 3))
    assert structural == ContextParagraph(
        ContextReason.STRUCTURAL_KIND,
        ContextRole.HEADING,
    )
    non_unique = _observation(_hit(
        {CompilerBlockKind.DIALOGUE_LABEL, CompilerBlockKind.PARAGRAPH},
        7,
        7,
    ))
    assert non_unique == ReviewParagraph(ReviewReason.NON_UNIQUE_CLAIMS)
    same_kind_twice = SourceBlockHit((
        BlockClaim(CompilerBlockKind.PARAGRAPH, ir.SourceSpan(7, 7), (0,)),
        BlockClaim(CompilerBlockKind.PARAGRAPH, ir.SourceSpan(7, 7), (1,)),
    ))
    assert _observation(same_kind_twice) == ReviewParagraph(ReviewReason.NON_UNIQUE_CLAIMS)
    unknown = _observation(_hit({CompilerBlockKind.UNKNOWN}, 2, 2))
    assert unknown == ReviewParagraph(ReviewReason.UNKNOWN_KIND)
    merged = _observation(_hit({CompilerBlockKind.PARAGRAPH}, 5, 7))
    assert merged == ReviewParagraph(ReviewReason.MERGED_BODY)


def test_structure_observation_distinguishes_dropped_from_unmapped(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    adapted = ir.Document(blocks=[
        ir.Paragraph(inlines=[ir.Text("prose")], source_span=ir.SourceSpan(0, 0)),
        ir.Heading(level=1, inlines=[ir.Text("H")], source_span=ir.SourceSpan(1, 1)),
        ir.Paragraph(inlines=[ir.Text("dropped")], source_span=ir.SourceSpan(2, 2)),
    ])
    seam = ir.Document(blocks=[
        adapted.blocks[0],
        adapted.blocks[1],
    ])
    source = _empty_source()
    monkeypatch.setattr(
        type(source),
        "content_ordinals",
        property(lambda _source: frozenset(range(4))),
    )
    monkeypatch.setattr("pancratius.docx_structure.da.adapt", lambda *_args: adapted)
    monkeypatch.setattr("pancratius.docx_structure.run", lambda *_args, **_kwargs: seam)

    observation = observe_structure(source, lang="ru")

    assert observation.source is source
    assert observation.lang == "ru"
    by_ordinal = observation.by_ordinal
    assert by_ordinal[docx_source.ParagraphOrdinal(0)] == BodyParagraph()
    assert by_ordinal[docx_source.ParagraphOrdinal(1)] == ContextParagraph(
        ContextReason.STRUCTURAL_KIND,
        ContextRole.HEADING,
    )
    assert by_ordinal[docx_source.ParagraphOrdinal(2)] == ContextParagraph(
        ContextReason.DROPPED_BY_STRUCTURAL_PIPELINE
    )
    assert by_ordinal[docx_source.ParagraphOrdinal(3)] == ReviewParagraph(
        ReviewReason.UNMAPPED_AT_ADAPTER
    )


def test_source_block_hits_projects_container_role_to_nested_members() -> None:
    paragraph = ir.Paragraph(
        inlines=[ir.Text("item")],
        source_span=ir.SourceSpan(7, 7),
    )
    blocks: list[ir.Block] = [ir.ListBlock(ordered=False, items=[[paragraph]])]

    hit = source_block_hits(blocks, {7})[7]

    assert hit.kinds == {CompilerBlockKind.LIST}
    assert hit.claims == (
        BlockClaim(
            kind=CompilerBlockKind.PARAGRAPH,
            span=ir.SourceSpan(7, 7),
            path=(0, 0, 0),
            context=CompilerBlockKind.LIST,
        ),
    )


def test_fold_result_uses_flow_claims_without_erasing_claim_cardinality() -> None:
    span = ir.SourceSpan(4, 4)
    label = BlockClaim(CompilerBlockKind.DIALOGUE_LABEL, span, (0,))
    paragraph = BlockClaim(CompilerBlockKind.PARAGRAPH, span, (1,))

    assert _fold_result(SourceBlockHit((label, paragraph))) == FoldDecision(
        FoldDisposition.FLOWING,
        (label, paragraph),
    )
    assert _fold_result(SourceBlockHit((paragraph, paragraph))) == FoldDecision(
        FoldDisposition.FLOWING,
        (paragraph, paragraph),
    )


def test_fold_result_keeps_mixed_flow_claims_as_a_typed_conflict() -> None:
    span = ir.SourceSpan(4, 4)
    paragraph = BlockClaim(CompilerBlockKind.PARAGRAPH, span, (0,))
    lineated = BlockClaim(CompilerBlockKind.LINEATED, span, (1,))

    assert _fold_result(SourceBlockHit((paragraph, lineated))) == FoldConflict(
        (paragraph, lineated)
    )
    observation = docx_structure.FoldObservation((
        (
            docx_source.ParagraphOrdinal(4),
            FoldConflict((paragraph, lineated)),
        ),
    ))
    assert observation.decisions == ()


def test_fold_observer_stops_before_register_assignment(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = _empty_source()
    seen: list[tuple[str | None, Context]] = []

    monkeypatch.setattr(
        docx_structure.da,
        "adapt",
        lambda _source, _media, _diagnostics: ir.Document(blocks=[]),
    )

    def fake_run(doc: ir.Document, context: Context, *, until: str | None = None) -> ir.Document:
        seen.append((until, context))
        return doc

    monkeypatch.setattr(docx_structure, "run", fake_run)

    assert docx_structure.observe_fold(source, lang="ru").entries == ()
    assert seen and seen[0][0] == POST_FOLD_SEAM
    assert not seen[0][1].scripture.by_ordinal


@pandoc_required
def test_semantic_surfaces_do_not_expand_enclosing_span_over_pagination(
    tmp_path: Path,
) -> None:
    """A folded block may enclose a skipped raw ordinal, but the ordinal is not a
    semantic contributor and must stay absent from every per-paragraph surface."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    document = Document()
    document.add_heading("Psalm", level=2)
    document.add_paragraph("My quiet light,")
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("burns in the heart.")
    path = tmp_path / "semantic-hole.docx"
    document.save(str(path))

    source = docx_source.read(path)
    assert source.paragraphs[2].disposition is docx_source.ParagraphDisposition.PAGINATION_ONLY
    classifications = docx_inspect.classify_blocks(source).by_source
    decisions = fold_decisions(source, lang="ru", apply_overrides=False)
    observation = observe_structure(source, lang="ru")

    assert classifications[1].span == ir.SourceSpan(1, 3)
    assert 2 not in classifications
    assert decisions[1] is decisions[3] is True
    assert 2 not in decisions
    assert 2 not in dict(observation.entries)


@pandoc_required
def test_fold_decisions_per_ordinal_surface(tmp_path: Path) -> None:
    """The per-`w:p`-ordinal prose/lineated surface the lineation gold joins on:
    prose stays False, an authored hard break is lineated regardless of Q2 register,
    structure is absent, and a folded
    couplet after a heading is True."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "Это длинное прозаическое предложение, которое заведомо длиннее любой "
        "стихотворной строки и читается как обычный абзац без всякой лиричности."
    )
    broken = doc.add_paragraph()
    run = broken.add_run("1. Вода")
    run.add_break()
    broken.add_run("Мир — как река. " * 12)
    doc.add_heading("Псалом", level=2)
    doc.add_paragraph("Свет мой тихий,")
    doc.add_paragraph("в сердце горит.")
    path = tmp_path / "fixture.docx"
    doc.save(str(path))

    decisions = fold_decisions(docx_source.read(path), lang="ru")

    assert decisions[0] is False
    assert decisions[1] is True   # the compiler emitted one lineated block; Q2 is separate
    assert 2 not in decisions     # the heading is structure, not a votable body line
    assert decisions[3] is True and decisions[4] is True


@pandoc_required
def test_fold_decisions_cover_register_quote_members(tmp_path: Path) -> None:
    """Register wrapping keeps fold coverage without conflating rendered hard lines:
    both quote members remain Paragraphs, so both are flowing at the fold seam. The
    hard break still renders as two-space lineation, independently test-pinned by
    `test_quote_member_hard_breaks_become_display_lines`."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph

    def set_border(paragraph: DocxParagraph, *sides: str) -> None:
        ppr = paragraph._p.get_or_add_pPr()  # fixture-only OOXML poke
        pbdr = ppr.makeelement(qn("w:pBdr"), {})
        for side in sides:
            el = ppr.makeelement(qn(f"w:{side}"), {qn("w:val"): "single", qn("w:sz"): "4"})
            pbdr.append(el)
        ppr.append(pbdr)

    doc = Document()
    filler = (
        "Это длинное прозаическое предложение, которое заведомо длиннее любой "
        "стихотворной строки и читается как обычный абзац без всякой лиричности."
    )
    for _ in range(8):
        doc.add_paragraph(filler)
    boxed = doc.add_paragraph(
        "7 Се, грядет с облаками, и узрит Его всякое око и те, которые пронзили Его."
    )
    set_border(boxed, "top", "bottom", "left", "right")
    ruled = doc.add_paragraph()
    run = ruled.add_run("Я — не форма,")
    run.add_break()
    ruled.add_run("но во всех формах живу.")
    set_border(ruled, "left")
    path = tmp_path / "fixture-borders.docx"
    doc.save(str(path))

    decisions = fold_decisions(docx_source.read(path), lang="ru")

    assert decisions[8] is False   # boxed prose verse: covered, prose register
    assert decisions[9] is False   # rendered hard lines, but never a LineatedBlock fold


@pandoc_required
def test_fold_decisions_en_edition_mirrors_ru(tmp_path: Path) -> None:
    """The EN editions get the same per-ordinal surface: EN prose stays False,
    an EN speaker turn (`Answer from the Creator:`) is structure — absent, never
    a verse line — and an EN couplet after a heading folds True."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "This is a long prose sentence that is obviously longer than any line of "
        "verse and reads like an ordinary paragraph without any lyricism at all."
    )
    turn = doc.add_paragraph()
    turn.add_run("Answer from the Creator:").bold = True
    doc.add_paragraph("A plain single answer sentence follows the speaker label.")
    doc.add_heading("Psalm", level=2)
    doc.add_paragraph("My quiet light,")
    doc.add_paragraph("burns in the heart.")
    path = tmp_path / "fixture-en.docx"
    doc.save(str(path))

    decisions = fold_decisions(docx_source.read(path), lang="en")

    assert decisions[0] is False
    assert decisions.get(1) is not True   # the speaker turn is never lineated
    assert decisions[2] is False
    assert 3 not in decisions             # the heading is structure
    assert decisions[4] is True and decisions[5] is True
