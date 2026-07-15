"""The semantics-preserving DOCX projection consumed by Pandoc.

Some source DOCX bind the correct OOXML drawing URIs to GENERIC prefixes
(`ns3:`/`ns5:`/`ns7:` …); pandoc 3.x then resolves no images and drops every one. The adapter
re-prefixes such a doc before pandoc reads it — changing prefixes only, never URIs, so the
recovered images are real and no text is lost. A conventionally-prefixed DOCX is passed through
untouched when it also has no page/column breaks.
"""

from __future__ import annotations

import shutil
import warnings
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

from pancratius import docx_adapter as da
from pancratius import docx_pandoc, docx_source, ir, ooxml
from pancratius.ir.inlines import inline_plain

pandoc_required = pytest.mark.skipif(
    shutil.which("pandoc") is None,
    reason="pandoc is required for importer-backed DOCX paths",
)

_REPO = Path(__file__).resolve().parents[2]
# A conventionally-prefixed book (images already work) and a generic-prefix book (images dropped
# by pandoc until canonicalized) — both small, both committed.
_CONVENTIONAL = _REPO / "src/content/books/11-kniga-ischezayushchego-ya/en.docx"
_GENERIC = _REPO / "src/content/books/27-mikki-17/ru.docx"
_BOOK17 = _REPO / "src/content/books/17-obolgannye-iisus-i-tvorets/ru.docx"
_BOOK39 = (
    _REPO
    / "src/content/books/39-kogda-bog-zagovoril-iz-mashiny-istoriya-iskina/ru.docx"
)
_GENERIC_IMAGE_RICH = (
    _REPO / "src/content/books/33-ya-esm-vsadnik-kon-i-mech/ru.docx"
)
_GENERIC_NAMESPACE_DOCX = (_GENERIC, _GENERIC_IMAGE_RICH)
_NO_STYLES = docx_source.ParagraphStyles()


def test_inline_text_keeps_break_policy_and_hides_opaque_payload() -> None:
    inlines = [
        {"t": "Str", "c": "first"},
        {"t": "SoftBreak"},
        {"t": "Str", "c": "second"},
        {"t": "LineBreak"},
        {
            "t": "Note",
            "c": [{"t": "Para", "c": [{"t": "Str", "c": "MINTED"}]}],
        },
        {"t": "Image", "c": [["", [], []], [], ["x.png", ""]]},
        {"t": "Cite", "c": [[], [{"t": "Str", "c": "cited"}]]},
        {"t": "Space"},
        {"t": "Math", "c": [{"t": "InlineMath"}, "x+y"]},
        {"t": "RawInline", "c": ["html", "<b>OPAQUE</b>"]},
        {"t": "Str", "c": "third"},
    ]

    assert (
        docx_pandoc.inline_text(
            inlines,
            soft_break=docx_pandoc.SoftBreakRendering.SPACE,
        )
        == "first second\ncited x+ythird"
    )
    assert (
        docx_pandoc.inline_text(
            inlines,
            soft_break=docx_pandoc.SoftBreakRendering.LINE,
        )
        == "first\nsecond\ncited x+ythird"
    )


@pytest.mark.parametrize(
    "inlines",
    [
        [{"t": "Str", "c": ["not text"]}],
        [{"t": "Cite", "c": [[], "not inlines"]}],
        [{"t": "Future", "c": [{"t": "Str", "c": "LEAK"}]}],
    ],
)
def test_inline_text_rejects_malformed_or_unknown_payload(inlines: object) -> None:
    with pytest.raises(docx_pandoc.PandocInlineError):
        docx_pandoc.inline_text(
            inlines,
            soft_break=docx_pandoc.SoftBreakRendering.SPACE,
        )


def test_conventional_docx_gains_only_source_anchors(tmp_path: Path) -> None:
    source = docx_source.read(_CONVENTIONAL)
    projected = docx_pandoc.project_package(source, tmp_path)
    assert projected != _CONVENTIONAL
    with zipfile.ZipFile(projected) as archive:
        root = ET.fromstring(archive.read(docx_source.DOCUMENT_PART))
    names = {
        element.get(f"{docx_source.W}name")
        for element in root.findall(f".//{docx_source.W}bookmarkStart")
    }
    expected = {
        docx_pandoc.source_anchor_name(ordinal)
        for ordinal in docx_pandoc.anchored_ordinals(source)
    }
    assert expected and expected <= names


def test_alphabetic_source_alias_is_canonicalized_for_pandoc() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" xmlns:foo="{ooxml.WP_NS}">'
        '<w:body><foo:inline/></w:body></w:document>'
    ).encode()

    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )

    assert b"<wp:inline" in projection
    assert ET.fromstring(projection).find(f".//{{{ooxml.WP_NS}}}inline") is not None


def _body_words(docx: Path, media: Path) -> list[str]:
    from pancratius import ir
    from pancratius.ir.inlines import inline_plain

    doc = da.adapt(docx_source.read(docx), media, [])
    text = " ".join(inline_plain(b.inlines) for b in doc.blocks
                    if isinstance(b, ir.Paragraph) and b.inlines)
    return text.split()


@pandoc_required
def test_canonicalization_loses_no_body_text(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Real comparison: import WITHOUT the rewrite (force passthrough) vs WITH it. The rewrite
    # is strictly ADDITIVE — it recovers images and the body words pandoc dropped alongside them
    # (e.g. textbox text), and removes nothing. So every passthrough word survives in the
    # rewritten import; a regression that corrupted existing text would drop one here.
    rewritten = _body_words(_GENERIC_IMAGE_RICH, tmp_path / "with")

    monkeypatch.setattr(
        docx_pandoc,
        "project_package",
        lambda source, _work_dir: source.path,
    )
    passthrough = _body_words(_GENERIC_IMAGE_RICH, tmp_path / "without")

    remaining = iter(rewritten)
    assert len(passthrough) > 1000
    assert all(any(candidate == word for candidate in remaining) for word in passthrough)
    assert len(rewritten) > len(passthrough)


def _break_docx(tmp_path: Path) -> Path:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("before")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("middle")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("line")
    paragraph.add_run().add_break(WD_BREAK.COLUMN)
    paragraph.add_run("after")
    path = tmp_path / "typed-breaks.docx"
    document.save(str(path))
    return path


def _pagination_only_docx(tmp_path: Path) -> Path:
    document = Document()
    document.add_paragraph("before")
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("after")
    page_before = document.add_paragraph()
    page_before.paragraph_format.page_break_before = True
    document.add_paragraph("end")
    document.add_paragraph()  # a real empty paragraph remains a structural boundary
    path = tmp_path / "pagination-only.docx"
    document.save(str(path))
    return path


def _fallback_pagination_only_docx(tmp_path: Path) -> Path:
    document = Document()
    document.add_paragraph("before")
    run = document.add_paragraph().add_run()
    run._r.append(
        parse_xml(
            f'<mc:AlternateContent xmlns:mc="{ooxml.MC_NS}" '
            f'xmlns:w="{ooxml.W_NS}" xmlns:x="urn:unsupported">'
            '<mc:Choice Requires="x"><w:t>INACTIVE</w:t></mc:Choice>'
            '<mc:Fallback><w:br w:type="page"/></mc:Fallback>'
            '</mc:AlternateContent>'
        )
    )
    document.add_paragraph("after")
    path = tmp_path / "fallback-pagination-only.docx"
    document.save(str(path))
    return path


def _opaque_pagination_docx(tmp_path: Path) -> Path:
    document = Document()
    document.add_paragraph("before")

    reference = document.add_paragraph()
    reference.add_run().add_break(WD_BREAK.PAGE)
    footnote = OxmlElement("w:footnoteReference")
    footnote.set(qn("w:id"), "1")
    reference.add_run()._r.append(footnote)

    equation = document.add_paragraph()
    equation.add_run().add_break(WD_BREAK.PAGE)
    equation._p.append(OxmlElement("m:oMath"))

    anchored = document.add_paragraph()
    anchored.add_run().add_break(WD_BREAK.PAGE)
    bookmark = OxmlElement("w:bookmarkStart")
    bookmark.set(qn("w:id"), "9")
    bookmark.set(qn("w:name"), "anchor")
    anchored._p.append(bookmark)
    anchored.add_run()._r.append(OxmlElement("w:fldChar"))

    section = document.add_paragraph()
    section.paragraph_format.page_break_before = True
    section._p.get_or_add_pPr().append(OxmlElement("w:sectPr"))
    document.add_paragraph("after")
    path = tmp_path / "opaque-pagination.docx"
    document.save(str(path))
    return path


def test_break_projection_neutralizes_only_pagination(tmp_path: Path) -> None:
    path = _break_docx(tmp_path)
    source = docx_source.read(path)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")

    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        document_xml,
        source,
        styles=source.styles,
    )

    root = ET.fromstring(projection)
    remaining = root.findall(f".//{docx_source.W}br")
    assert len(remaining) == 1
    assert docx_source.BreakKind.from_ooxml(
        remaining[0].get(f"{docx_source.W}type")
    ) is docx_source.BreakKind.LINE


def test_inactive_choice_paragraph_is_never_removed() -> None:
    xml = (
        f'<w:footnotes xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
        '<w:footnote w:id="1"><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:p><w:r>'
        '<w:br w:type="page"/>'
        '</w:r></w:p></mc:Choice>'
        '<mc:Fallback><w:t>BASELINE</w:t></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:footnote></w:footnotes>'
    ).encode()

    projection = docx_pandoc._project_part(
        "word/footnotes.xml",
        xml,
        styles=_NO_STYLES,
    )

    root = ET.fromstring(projection)
    paragraphs = root.findall(f".//{docx_source.W}p")
    assert len(paragraphs) == 1
    assert root.find(f".//{docx_source.W}br") is None
    assert root.find(f".//{ooxml.MC_ALTERNATE_CONTENT}") is None
    assert "".join(element.text or "" for element in root.iter(f"{docx_source.W}t")) == "BASELINE"


@pandoc_required
def test_pagination_only_paragraphs_do_not_mint_empty_ir_blocks(tmp_path: Path) -> None:
    path = _pagination_only_docx(tmp_path)
    source = docx_source.read(path)

    assert [
        paragraph.reconciliation_position.value
        if paragraph.reconciliation_position is not None
        else None
        for paragraph in source.paragraphs
    ] == [0, None, 1, None, 2, 3]
    assert source.paragraphs[1].disposition is docx_source.ParagraphDisposition.PAGINATION_ONLY
    assert source.paragraphs[3].disposition is docx_source.ParagraphDisposition.PAGINATION_ONLY

    projected = docx_pandoc.project_package(source, tmp_path / "projected")
    with zipfile.ZipFile(projected) as archive:
        root = ET.fromstring(archive.read(docx_source.DOCUMENT_PART))
    body = root.find(f"{docx_source.W}body")
    assert body is not None
    assert len(docx_source.body_paragraph_elements(body)) == 4 + 1  # + the anchor farm

    imported = da.adapt(source, tmp_path / "media", [])
    assert [
        inline_plain(block.inlines)
        for block in imported.blocks
        if isinstance(block, ir.Paragraph)
    ] == ["before", "after", "end", ""]


def test_pagination_never_deletes_opaque_source_payload(tmp_path: Path) -> None:
    path = _opaque_pagination_docx(tmp_path)
    source = docx_source.read(path)
    assert [paragraph.disposition for paragraph in source.paragraphs[1:5]] == [
        docx_source.ParagraphDisposition.NON_TEXT,
        docx_source.ParagraphDisposition.NON_TEXT,
        docx_source.ParagraphDisposition.NON_TEXT,
        docx_source.ParagraphDisposition.NON_TEXT,
    ]

    projected = docx_pandoc.project_package(source, tmp_path / "projected")
    with zipfile.ZipFile(projected) as archive:
        root = ET.fromstring(archive.read(docx_source.DOCUMENT_PART))

    assert len(root.findall(f".//{docx_source.W}footnoteReference")) == 1
    assert len(root.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")) == 1
    foreign_bookmarks = [
        element
        for element in root.findall(f".//{docx_source.W}bookmarkStart")
        if not (element.get(f"{docx_source.W}name") or "").startswith(
            docx_pandoc.SOURCE_ANCHOR_PREFIX
        )
    ]
    assert len(foreign_bookmarks) == 1
    assert len(root.findall(f".//{docx_source.W}fldChar")) == 1
    assert len(root.findall(f".//{docx_source.W}sectPr")) >= 2


def test_alternate_content_uses_fallback_for_source_and_pagination() -> None:
    """An unselected Choice cannot contribute text or pagination evidence."""
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
        '<w:body><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:r><w:br w:type="page"/></w:r></mc:Choice>'
        '<mc:Fallback><w:br/><w:t>VISIBLE FALLBACK</w:t></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:body></w:document>'
    ).encode()
    paragraph = ET.fromstring(xml).find(f".//{docx_source.W}p")
    assert paragraph is not None

    semantics = docx_source.analyze_paragraph(
        paragraph,
        styles=docx_source.ParagraphStyles(),
    )
    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )

    assert semantics.content.reading.strip() == "VISIBLE FALLBACK"
    assert semantics.content.lineated == "\nVISIBLE FALLBACK"
    assert semantics.content.breaks == (docx_source.BreakKind.LINE,)
    assert semantics.disposition is docx_source.ParagraphDisposition.CONTENT
    assert b"VISIBLE FALLBACK" in projection
    assert ET.fromstring(projection).find(f".//{ooxml.MC_ALTERNATE_CONTENT}") is None


def test_natural_lines_normalize_layout_tabs_like_reading_text() -> None:
    content = docx_source.ParagraphContent(
        (
            docx_source.TextAtom("chapter\t\u00a012"),
            docx_source.BreakKind.LINE,
            docx_source.TextAtom("next  line"),
        )
    )

    assert content.reading == "chapter 12 next line"
    assert content.line_segments == ("chapter 12", "next line")


def test_story_contents_share_the_canonical_break_and_compatibility_grammar() -> None:
    root = ET.fromstring(
        f'<w:footnotes xmlns:w="{ooxml.W_NS}" xmlns:mc="{ooxml.MC_NS}">'
        '<w:footnote><w:p><w:r><w:t>one</w:t><w:br w:type="page"/>'
        '<mc:AlternateContent><mc:Choice Requires="x"><w:t>hidden</w:t></mc:Choice>'
        '<mc:Fallback><w:br/><w:t>two</w:t></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:footnote></w:footnotes>'
    )

    content, = docx_source.story_contents(root)

    assert content.reading == "one two"
    assert content.line_segments == ("one", "two")
    assert content.breaks == (docx_source.BreakKind.PAGE, docx_source.BreakKind.LINE)


def _add_story(path: Path, part: docx_source.StoryPart, body: str) -> None:
    root = part.name.lower()
    payload = f'<w:{root} xmlns:w="{ooxml.W_NS}" xmlns:mc="{ooxml.MC_NS}">{body}</w:{root}>'
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr(part.value, payload.encode())


def test_read_story_handles_required_and_optional_parts(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("body")
    path = tmp_path / "stories.docx"
    document.save(str(path))

    assert tuple(content.reading for content in docx_source.read_story(
        path, docx_source.StoryPart.DOCUMENT
    )) == ("body",)
    assert docx_source.read_story(path, docx_source.StoryPart.FOOTNOTES) == ()
    assert docx_source.read_story(path, docx_source.StoryPart.ENDNOTES) == ()

    empty = tmp_path / "missing-document.docx"
    with zipfile.ZipFile(empty, "w"):
        pass
    with pytest.raises(docx_source.DocxSourceError, match=r"missing word/document\.xml"):
        docx_source.read_story(empty, docx_source.StoryPart.DOCUMENT)


@pytest.mark.parametrize(
    "part",
    [docx_source.StoryPart.FOOTNOTES, docx_source.StoryPart.ENDNOTES],
)
def test_read_story_uses_typed_breaks_and_selected_fallback(
    tmp_path: Path,
    part: docx_source.StoryPart,
) -> None:
    document = Document()
    path = tmp_path / f"{part.name.lower()}.docx"
    document.save(str(path))
    _add_story(
        path,
        part,
        '<w:p><w:r><w:t>one</w:t><w:br w:type="page"/>'
        '<mc:AlternateContent><mc:Choice Requires="x"><w:t>hidden</w:t></mc:Choice>'
        '<mc:Fallback><w:br/><w:t>two</w:t></mc:Fallback>'
        '</mc:AlternateContent><w:br w:type="column"/></w:r></w:p>',
    )

    content, = docx_source.read_story(path, part)

    assert content.reading == "one two"
    assert content.line_segments == ("one", "two")
    assert content.breaks == (
        docx_source.BreakKind.PAGE,
        docx_source.BreakKind.LINE,
        docx_source.BreakKind.COLUMN,
    )


def test_read_story_reports_package_part_and_paragraph_for_unknown_break(
    tmp_path: Path,
) -> None:
    document = Document()
    path = tmp_path / "unknown-break.docx"
    document.save(str(path))
    _add_story(
        path,
        docx_source.StoryPart.FOOTNOTES,
        '<w:p><w:r><w:t>first</w:t></w:r></w:p>'
        '<w:p><w:r><w:br w:type="future-layout"/></w:r></w:p>',
    )

    with pytest.raises(
        docx_source.DocxSourceError,
        match=(
            r"unknown-break\.docx: cannot read word/footnotes\.xml: paragraph 1: "
            r"unsupported w:br type 'future-layout'"
        ),
    ):
        docx_source.read_story(path, docx_source.StoryPart.FOOTNOTES)


def test_alternate_content_without_fallback_remains_opaque() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
        '<w:body><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:t>CHOICE ONLY</w:t></mc:Choice>'
        '</mc:AlternateContent></w:r></w:p></w:body></w:document>'
    ).encode()
    paragraph = ET.fromstring(xml).find(f".//{docx_source.W}p")
    assert paragraph is not None

    semantics = docx_source.analyze_paragraph(
        paragraph,
        styles=docx_source.ParagraphStyles(),
    )
    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )

    assert semantics.content.atoms == ()
    assert semantics.has_opaque_payload
    assert semantics.disposition is docx_source.ParagraphDisposition.NON_TEXT
    assert len(ET.fromstring(projection).findall(f".//{docx_source.W}p")) == 1
    assert b"CHOICE ONLY" not in projection


def test_pagination_projection_preserves_numbered_paragraph_semantics() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}"><w:body><w:p>'
        '<w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>'
        '<w:r><w:br w:type="page"/></w:r>'
        '</w:p></w:body></w:document>'
    ).encode()
    paragraph = ET.fromstring(xml).find(f".//{docx_source.W}p")
    assert paragraph is not None

    semantics = docx_source.analyze_paragraph(
        paragraph,
        styles=docx_source.ParagraphStyles(),
    )
    projected = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )

    assert semantics.payload.kinds == frozenset(
        {docx_source.ParagraphPayloadKind.DIRECT_NUMBERING}
    )
    assert semantics.disposition is docx_source.ParagraphDisposition.NON_TEXT
    assert ET.fromstring(projected).find(f".//{docx_source.W}numPr") is not None


def test_pagination_projection_resolves_numbering_but_not_ordinary_styles(
    tmp_path: Path,
) -> None:
    document = Document()
    ordinary = document.add_paragraph(style="Quote")
    ordinary.add_run().add_break(WD_BREAK.PAGE)
    numbered_style = document.styles.add_style(
        "Derived Number",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    numbered_style.base_style = document.styles["List Number"]
    numbered = document.add_paragraph(style=numbered_style)
    numbered.add_run().add_break(WD_BREAK.PAGE)
    path = tmp_path / "styles.docx"
    document.save(str(path))

    source = docx_source.read(path)
    projected = docx_pandoc.project_package(source, tmp_path / "projected")
    with zipfile.ZipFile(projected) as archive:
        root = ET.fromstring(archive.read(docx_source.DOCUMENT_PART))

    assert [paragraph.disposition for paragraph in source.paragraphs] == [
        docx_source.ParagraphDisposition.PAGINATION_ONLY,
        docx_source.ParagraphDisposition.NON_TEXT,
    ]
    assert source.paragraphs[1].semantics.payload.kinds == frozenset(
        {docx_source.ParagraphPayloadKind.RESOLVED_NUMBERING}
    )
    remaining = root.findall(f".//{docx_source.W}p")
    assert len(remaining) == 1
    assert remaining[0].find(f"{docx_source.W}pPr/{docx_source.W}pStyle") is not None


@pandoc_required
def test_book39_styled_pagination_does_not_mint_an_empty_ir_block(
    tmp_path: Path,
) -> None:
    source = docx_source.read(_BOOK39)

    assert source.paragraphs[46].disposition is docx_source.ParagraphDisposition.PAGINATION_ONLY
    adapted = da.adapt(source, tmp_path / "media", [])

    assert not (
        isinstance(adapted.blocks[46], ir.Paragraph)
        and adapted.blocks[46].empty
        and adapted.blocks[46].source_span is None
    )


def test_alternate_content_outside_run_contributes_no_atoms() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}"><w:body><w:p><mc:AlternateContent>'
        '<mc:Fallback><w:r><w:t>WRONG PLACEMENT</w:t></w:r></mc:Fallback>'
        '</mc:AlternateContent></w:p></w:body></w:document>'
    ).encode()
    paragraph = ET.fromstring(xml).find(f".//{docx_source.W}p")
    assert paragraph is not None
    semantics = docx_source.analyze_paragraph(
        paragraph,
        styles=docx_source.ParagraphStyles(),
    )
    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )

    root = ET.fromstring(projection)
    assert semantics.content.atoms == ()
    assert semantics.disposition is docx_source.ParagraphDisposition.NON_TEXT
    assert len(root.findall(f".//{docx_source.W}p")) == 1
    assert root.find(f".//{ooxml.MC_ALTERNATE_CONTENT}") is None
    assert b"WRONG PLACEMENT" not in projection


def test_nested_alternate_content_keeps_its_original_capability_context() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
        '<w:body><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:t>OUTER CHOICE</w:t></mc:Choice>'
        '<mc:Fallback><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:t>INNER CHOICE</w:t></mc:Choice>'
        '<mc:Fallback><w:t>LEAKED</w:t></mc:Fallback>'
        '</mc:AlternateContent></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:body></w:document>'
    ).encode()
    paragraph = ET.fromstring(xml).find(f".//{docx_source.W}p")
    assert paragraph is not None

    semantics = docx_source.analyze_paragraph(
        paragraph,
        styles=docx_source.ParagraphStyles(),
    )
    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )
    projected_root = ET.fromstring(projection)

    assert semantics.content.atoms == ()
    assert semantics.disposition is docx_source.ParagraphDisposition.NON_TEXT
    assert projected_root.find(f".//{ooxml.MC_ALTERNATE_CONTENT}") is None
    assert b"LEAKED" not in projection


def test_removed_choices_cannot_conflict_in_namespace_closure() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" xmlns:mc="{ooxml.MC_NS}">'
        '<w:body><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice xmlns:x="urn:choice-one" Requires="x">'
        '<w:t>FIRST</w:t></mc:Choice>'
        '<mc:Choice xmlns:x="urn:choice-two" Requires="x">'
        '<w:t>SECOND</w:t></mc:Choice>'
        '<mc:Fallback><w:t>BASELINE</w:t></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:body></w:document>'
    ).encode()

    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )

    assert b"BASELINE" in projection
    assert b"FIRST" not in projection
    assert b"SECOND" not in projection


def test_removed_pagination_attributes_cannot_conflict_in_namespace_closure() -> None:
    xml = (
        f'<w:document xmlns:w="{ooxml.W_NS}" xmlns:mc="{ooxml.MC_NS}">'
        '<w:body>'
        '<w:p><w:r><w:t>FIRST</w:t>'
        '<w:br xmlns:x="urn:first" w:type="page" mc:Ignorable="x"/>'
        '</w:r></w:p>'
        '<w:p><w:r><w:t>SECOND</w:t>'
        '<w:br xmlns:x="urn:second" w:type="page" mc:Ignorable="x"/>'
        '</w:r></w:p>'
        '</w:body></w:document>'
    ).encode()

    projection = docx_pandoc._project_part(
        docx_source.DOCUMENT_PART,
        xml,
        styles=_NO_STYLES,
    )
    root = ET.fromstring(projection)

    assert [element.text for element in root.iter(f"{docx_source.W}t")] == [
        "FIRST",
        " ",
        "SECOND",
        " ",
    ]
    assert not ooxml.unresolved_prefix_value_references(projection)


def test_alternate_content_rejects_duplicate_fallback_with_context() -> None:
    xml = (
        f'<w:footnotes xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
        '<w:footnote w:id="1"><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:t>CHOICE</w:t></mc:Choice>'
        '<mc:Fallback><w:t>FIRST</w:t></mc:Fallback>'
        '<mc:Fallback><w:t>SECOND</w:t></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:footnote></w:footnotes>'
    ).encode()

    with pytest.raises(
        docx_source.DocxSourceError,
        match=(
            r"word/footnotes\.xml: paragraph 0: "
            r"mc:AlternateContent has multiple fallback branches"
        ),
    ):
        docx_pandoc._project_part(
            "word/footnotes.xml",
            xml,
            styles=_NO_STYLES,
        )


def test_only_selected_fallback_pagination_is_neutralized() -> None:
    xml = (
        f'<w:footnotes xmlns:w="{ooxml.W_NS}" '
        f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
        '<w:footnote w:id="1"><w:p><w:r><mc:AlternateContent>'
        '<mc:Choice Requires="x"><w:t>CHOICE</w:t></mc:Choice>'
        '<mc:Fallback><w:br w:type="page"/></mc:Fallback>'
        '</mc:AlternateContent></w:r></w:p></w:footnote></w:footnotes>'
    ).encode()

    projection = docx_pandoc._project_part(
        "word/footnotes.xml",
        xml,
        styles=_NO_STYLES,
    )
    root = ET.fromstring(projection)

    assert root.find(f".//{docx_source.W}br") is None
    assert root.find(f".//{docx_source.W}p") is None
    assert b"CHOICE" not in projection


@pandoc_required
def test_fallback_only_pagination_paragraph_mints_no_ir_block(tmp_path: Path) -> None:
    path = _fallback_pagination_only_docx(tmp_path)
    source = docx_source.read(path)

    assert [paragraph.disposition for paragraph in source.paragraphs] == [
        docx_source.ParagraphDisposition.CONTENT,
        docx_source.ParagraphDisposition.PAGINATION_ONLY,
        docx_source.ParagraphDisposition.CONTENT,
    ]
    projected = docx_pandoc.project_package(source, tmp_path / "projected")
    with zipfile.ZipFile(projected) as archive:
        root = ET.fromstring(archive.read(docx_source.DOCUMENT_PART))
    assert root.find(f".//{ooxml.MC_ALTERNATE_CONTENT}") is None
    assert len(root.findall(f".//{docx_source.W}p")) == 2 + 1  # + the anchor farm

    imported = da.adapt(source, tmp_path / "media", [])
    assert [
        inline_plain(block.inlines)
        for block in imported.blocks
        if isinstance(block, ir.Paragraph)
    ] == ["before", "after"]


def _alternate_content_docx(tmp_path: Path) -> Path:
    document = Document()
    run = document.add_paragraph().add_run()
    run._r.append(
        parse_xml(
            f'<mc:AlternateContent xmlns:mc="{ooxml.MC_NS}" '
            f'xmlns:w="{ooxml.W_NS}" '
            'xmlns:wps="http://schemas.microsoft.com/office/word/2010/'
            'wordprocessingShape" xmlns:x="urn:unsupported">'
            '<mc:Choice Requires="wps"><w:t>MODERN</w:t></mc:Choice>'
            '<mc:Choice Requires="x"><w:t>FUTURE</w:t></mc:Choice>'
            '<mc:Fallback><w:t>BASELINE</w:t></mc:Fallback>'
            '</mc:AlternateContent>'
        )
    )
    path = tmp_path / "alternate-content.docx"
    document.save(str(path))
    return path


@pandoc_required
def test_source_and_pandoc_share_the_mc_fallback_profile(tmp_path: Path) -> None:
    path = _alternate_content_docx(tmp_path)
    source = docx_source.read(path)
    imported = da.adapt(source, tmp_path / "media", [])
    paragraphs = [block for block in imported.blocks if isinstance(block, ir.Paragraph)]

    assert [paragraph.text for paragraph in source.paragraphs] == ["BASELINE"]
    assert [inline_plain(block.inlines) for block in paragraphs] == ["BASELINE"]


@pytest.mark.parametrize("part", ["word/footnotes.xml", "word/endnotes.xml"])
def test_note_story_parts_neutralize_pagination_without_modeling_notes(part: str) -> None:
    root_name = "footnotes" if "footnotes" in part else "endnotes"
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        f'<w:{root_name} xmlns:w="{ooxml.W_NS}">'
        '<w:p><w:r><w:t>before</w:t><w:br w:type="page"/>'
        '<w:t>after</w:t><w:br/></w:r></w:p>'
        f'</w:{root_name}>'
    ).encode()

    projection = docx_pandoc._project_part(part, xml, styles=_NO_STYLES)

    rendered = ET.fromstring(projection)
    breaks = rendered.findall(f".//{docx_source.W}br")
    assert len(breaks) == 1
    assert breaks[0].get(f"{docx_source.W}type") is None


def test_story_projection_reports_unknown_break_with_part_and_paragraph() -> None:
    xml = (
        f'<w:footnotes xmlns:w="{ooxml.W_NS}">'
        '<w:footnote w:id="1"><w:p><w:r>'
        '<w:br w:type="future-layout"/>'
        '</w:r></w:p></w:footnote></w:footnotes>'
    ).encode()

    with pytest.raises(
        docx_source.DocxSourceError,
        match=(
            r"word/footnotes\.xml: paragraph 0: "
            r"unsupported w:br type 'future-layout'"
        ),
    ):
        docx_pandoc._project_part(
            "word/footnotes.xml",
            xml,
            styles=_NO_STYLES,
        )


def test_break_validation_applies_only_to_selected_fallback() -> None:
    def story(choice_type: str, fallback_type: str) -> bytes:
        return (
            f'<w:footnotes xmlns:w="{ooxml.W_NS}" '
            f'xmlns:mc="{ooxml.MC_NS}" xmlns:x="urn:unsupported">'
            '<w:footnote w:id="1"><w:p><w:r><mc:AlternateContent>'
            f'<mc:Choice Requires="x"><w:br w:type="{choice_type}"/></mc:Choice>'
            f'<mc:Fallback><w:br w:type="{fallback_type}"/></mc:Fallback>'
            '</mc:AlternateContent></w:r></w:p></w:footnote></w:footnotes>'
        ).encode()

    projection = docx_pandoc._project_part(
        "word/footnotes.xml",
        story("future-layout", "page"),
        styles=_NO_STYLES,
    )
    assert ET.fromstring(projection).find(f".//{docx_source.W}p") is None

    with pytest.raises(
        docx_source.DocxSourceError,
        match=(
            r"word/footnotes\.xml: paragraph 0: "
            r"unsupported w:br type 'future-layout'"
        ),
    ):
        docx_pandoc._project_part(
            "word/footnotes.xml",
            story("page", "future-layout"),
            styles=_NO_STYLES,
        )


def test_note_story_parts_remove_only_payload_free_pagination_paragraphs() -> None:
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        f'<w:footnotes xmlns:w="{ooxml.W_NS}">'
        '<w:footnote w:id="1"><w:p><w:r><w:br w:type="page"/></w:r></w:p></w:footnote>'
        '<w:footnote w:id="2"><w:p><w:r><w:br w:type="page"/>'
        '<w:footnoteReference w:id="2"/></w:r></w:p></w:footnote>'
        '</w:footnotes>'
    ).encode()

    projection = docx_pandoc._project_part(
        "word/footnotes.xml",
        xml,
        styles=_NO_STYLES,
    )

    root = ET.fromstring(projection)
    assert len(root.findall(f".//{docx_source.W}p")) == 1
    assert len(root.findall(f".//{docx_source.W}footnoteReference")) == 1


def test_disabled_page_break_before_remains_a_structural_empty_paragraph() -> None:
    xml = (
        f'<w:footnotes xmlns:w="{ooxml.W_NS}"><w:footnote w:id="1">'
        '<w:p><w:pPr><w:pageBreakBefore w:val="off"/></w:pPr></w:p>'
        '</w:footnote></w:footnotes>'
    ).encode()
    paragraph = ET.fromstring(xml).find(f".//{docx_source.W}p")
    assert paragraph is not None

    semantics = docx_source.analyze_paragraph(
        paragraph,
        styles=docx_source.ParagraphStyles(),
    )
    projection = docx_pandoc._project_part(
        "word/footnotes.xml",
        xml,
        styles=_NO_STYLES,
    )

    assert not semantics.page_break_before
    assert semantics.disposition is docx_source.ParagraphDisposition.STRUCTURAL_EMPTY
    assert len(ET.fromstring(projection).findall(f".//{docx_source.W}p")) == 1


@pytest.mark.parametrize("path", [_BOOK17, _GENERIC], ids=["book17", "book27"])
def test_projected_parts_close_every_lexical_namespace_reference(
    path: Path,
    tmp_path: Path,
) -> None:
    projected = docx_pandoc.project_package(docx_source.read(path), tmp_path)
    with zipfile.ZipFile(projected) as archive:
        for part in docx_pandoc._PANDOC_STORY_PARTS:
            if part in archive.namelist():
                assert ooxml.unresolved_prefix_value_references(archive.read(part)) == set()


@pytest.mark.parametrize(
    "path",
    _GENERIC_NAMESPACE_DOCX,
    ids=lambda path: path.parent.name.split("-", 1)[0],
)
def test_namespace_projection_preserves_lexical_meaning_across_generic_books(
    path: Path,
) -> None:
    """All known generic-prefix sources retain scoped MC namespace meaning."""
    with zipfile.ZipFile(path) as archive:
        source_xml = archive.read(docx_source.DOCUMENT_PART)
    source = ooxml.parse_xml(source_xml)
    projected = ooxml.serialize_xml(
        source,
        bindings=docx_pandoc._PANDOC_NAMESPACE_BINDINGS,
    )

    assert ET.tostring(ET.fromstring(projected)) == ET.tostring(source.root)
    common = {binding.prefix: binding.uri for binding in ooxml.COMMON_NAMESPACES}
    expected = tuple(
        ooxml.PrefixValueReference(
            reference.prefix,
            reference.uri or common[reference.prefix],
        )
        for reference in ooxml.prefix_value_references_in_xml(source_xml)
    )
    assert ooxml.prefix_value_references_in_xml(projected) == expected


def _zip_fingerprint(info: zipfile.ZipInfo) -> tuple[object, ...]:
    return (
        info.filename,
        info.date_time,
        info.compress_type,
        info.comment,
        info.extra,
        info.internal_attr,
        info.external_attr,
        info.create_system,
        info.flag_bits,
        info.create_version,
        info.extract_version,
        info.volume,
        info.reserved,
    )


def test_projected_package_preserves_entry_order_metadata_and_duplicates(
    tmp_path: Path,
) -> None:
    path = _break_docx(tmp_path)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with zipfile.ZipFile(path, "a") as archive:
            archive.comment = b"preserve package comment"
            archive.writestr("custom/duplicate.bin", b"first")
            archive.writestr("custom/duplicate.bin", b"second")
    source = docx_source.read(path)

    with zipfile.ZipFile(path) as archive:
        original_comment = archive.comment
        original_payloads = [archive.read(info) for info in archive.infolist()]
        before = [_zip_fingerprint(info) for info in archive.infolist()]
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        projected = docx_pandoc.project_package(source, tmp_path / "projected")
    with zipfile.ZipFile(projected) as archive:
        assert archive.comment == original_comment
        after = [_zip_fingerprint(info) for info in archive.infolist()]
        duplicate_payloads = [
            archive.read(info)
            for info in archive.infolist()
            if info.filename == "custom/duplicate.bin"
        ]
        projected_payloads = [archive.read(info) for info in archive.infolist()]
    assert after == before
    assert duplicate_payloads == [b"first", b"second"]
    for index, (original, projected) in enumerate(
        zip(original_payloads, projected_payloads, strict=True)
    ):
        if before[index][0] != docx_source.DOCUMENT_PART:
            assert projected == original


def test_projection_rejects_duplicate_document_story_part(tmp_path: Path) -> None:
    path = _break_docx(tmp_path)
    with zipfile.ZipFile(path) as archive:
        duplicate = archive.read(docx_source.DOCUMENT_PART)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with zipfile.ZipFile(path, "a") as archive:
            archive.writestr(docx_source.DOCUMENT_PART, duplicate)
    source = docx_source.read(path)

    with pytest.raises(
        RuntimeError,
        match=r"duplicate Word story part.*word/document\.xml",
    ):
        docx_pandoc.project_package(source, tmp_path / "projected")


def test_projection_fails_closed_when_an_optional_story_part_is_malformed(
    tmp_path: Path,
) -> None:
    path = _break_docx(tmp_path)
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", b"<w:footnotes")
    source = docx_source.read(path)

    with pytest.raises(RuntimeError, match=r"word/footnotes\.xml"):
        docx_pandoc.project_package(source, tmp_path / "projected")


@pandoc_required
def test_pandoc_projection_keeps_line_breaks_and_demotes_pagination(
    tmp_path: Path,
) -> None:
    path = _break_docx(tmp_path)
    document = da.adapt(docx_source.read(path), tmp_path / "media", [])
    paragraph = next(block for block in document.blocks if isinstance(block, ir.Paragraph))

    assert inline_plain(paragraph.inlines) == "before middle line after"
    assert sum(isinstance(inline, ir.LineBreak) for inline in paragraph.inlines) == 1


@pandoc_required
def test_book17_former_sidecar_is_output_equivalent(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    from pancratius import docx_conversion

    def convert(media: str) -> docx_conversion.ConvertedDocx:
        return docx_conversion.convert_single_docx(
            _BOOK17,
            kind="book",
            lang="ru",
            work_key="book:17",
            title="Оболганные Иисус и Творец",
            title_index={},
            media_out=tmp_path / media,
        )

    monkeypatch.setattr(
        docx_conversion.lineation_overrides,
        "load_overrides",
        lambda _source: {140: "prose", 141: "prose"},
    )
    pinned = convert("pinned")
    monkeypatch.setattr(
        docx_conversion.lineation_overrides,
        "load_overrides",
        lambda _source: {},
    )
    unpinned = convert("unpinned")

    assert (unpinned.body, unpinned.bibliography, unpinned.cross_refs) == (
        pinned.body,
        pinned.bibliography,
        pinned.cross_refs,
    )

# ---------------------------------------------------------------------------
# Source anchors: exact leaf provenance through Pandoc
# ---------------------------------------------------------------------------


def _provenance_message(diagnostics: list[ir.Diagnostic]) -> str:
    return next(d.message for d in diagnostics if d.code == "import.provenance")


def test_anchors_give_exact_identity_to_duplicate_and_nested_paragraphs(
    tmp_path: Path,
) -> None:
    """The two shapes content matching cannot resolve: duplicate text, and a
    paragraph nested inside a container. Anchors identify both exactly."""
    document = Document()
    document.add_paragraph("Repeated dedication")            # ordinal 0
    document.add_paragraph("Repeated dedication", style="Quote")  # ordinal 1, quote member
    document.add_paragraph()                                  # ordinal 2, empty
    document.add_paragraph("***")                             # ordinal 3, punctuation-only
    path = tmp_path / "anchors.docx"
    document.save(str(path))

    source = docx_source.read(path)
    diagnostics: list[ir.Diagnostic] = []
    imported = da.adapt(source, tmp_path / "media", diagnostics)

    quote = next(b for b in imported.blocks if isinstance(b, ir.QuoteBlock))
    (member,) = [m for m in quote.blocks if isinstance(m, ir.Paragraph)]
    assert member.source_span is not None and (
        member.source_span.start, member.source_span.end
    ) == (1, 1)
    assert quote.source_span == member.source_span
    top_paragraphs = [b for b in imported.blocks if isinstance(b, ir.Paragraph)]
    first = next(p for p in top_paragraphs if p.inlines and not p.empty)
    assert first.source_span is not None and (
        first.source_span.start, first.source_span.end
    ) == (0, 0)
    stars = next(
        p for p in top_paragraphs if p.inlines and ir.Text("***") in p.inlines
    )
    assert stars.source_span is not None and (
        stars.source_span.start, stars.source_span.end
    ) == (3, 3)
    assert "unclaimed-content=0" in _provenance_message(diagnostics)


def test_anchor_farm_never_reaches_the_ir(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("only content")
    path = tmp_path / "farm.docx"
    document.save(str(path))

    source = docx_source.read(path)
    imported = da.adapt(source, tmp_path / "media", [])

    texts = [
        "".join(t.value for t in b.inlines if isinstance(t, ir.Text))
        for b in imported.blocks
        if isinstance(b, ir.Paragraph)
    ]
    assert texts == ["only content"]
    assert not any(
        isinstance(i, ir.Link)
        for b in imported.blocks
        if isinstance(b, ir.Paragraph)
        for i in b.inlines
    )


def test_anchor_extraction_leaves_no_phantom_trailing_inline(tmp_path: Path) -> None:
    """The injected bookmark must not leave residue: a single-run bold paragraph
    stays a SINGLE inline (strong-opener detection depends on it), and a trailing
    hard break stays trimmed exactly as Pandoc trims it unanchored."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Потому запомни:").bold = True
    trailed = document.add_paragraph()
    trailed.add_run("title line")
    trailed.add_run().add_break(WD_BREAK.LINE)
    path = tmp_path / "residue.docx"
    document.save(str(path))

    source = docx_source.read(path)
    imported = da.adapt(source, tmp_path / "media", [])

    bold, titled = (b for b in imported.blocks if isinstance(b, ir.Paragraph))
    assert len(bold.inlines) == 1 and isinstance(bold.inlines[0], ir.Emphasis)
    assert not any(isinstance(i, ir.LineBreak) for i in titled.inlines)


def test_unclaimed_content_ordinals_fail_loud(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Provenance loss must never ship silently: dropping a claimed block from the
    AST surfaces the ordinal in the provenance warning."""
    document = Document()
    document.add_paragraph("kept")
    document.add_paragraph("lost")
    path = tmp_path / "unclaimed.docx"
    document.save(str(path))
    source = docx_source.read(path)

    real_run_json = docx_pandoc.run_json

    def dropping_run_json(
        src: docx_source.DocxSourceDocument, media: Path
    ) -> tuple[dict[str, object], str]:
        ast, warns = real_run_json(src, media)
        blocks = ast.get("blocks") or []
        kept = [b for b in blocks if "lost" not in str(b)]
        return {**ast, "blocks": kept}, warns

    monkeypatch.setattr(docx_pandoc, "run_json", dropping_run_json)
    diagnostics: list[ir.Diagnostic] = []
    da.adapt(source, tmp_path / "media", diagnostics)

    warning = next(
        d for d in diagnostics if d.code == "import.provenance-unclaimed"
    )
    assert "1 content paragraph" in warning.message


def test_heading_provenance_recovered_through_farm_rewrite(tmp_path: Path) -> None:
    """Pandoc folds a heading's bookmark into the Header id and rewrites the farm
    link; position-based recovery must hand the ordinal back to the Heading."""
    document = Document()
    document.add_heading("Chapter of Light", level=1)   # ordinal 0
    document.add_paragraph("Body under the heading.")   # ordinal 1
    path = tmp_path / "heading.docx"
    document.save(str(path))

    source = docx_source.read(path)
    diagnostics: list[ir.Diagnostic] = []
    imported = da.adapt(source, tmp_path / "media", diagnostics)

    heading = next(b for b in imported.blocks if isinstance(b, ir.Heading))
    assert heading.source_span is not None and (
        heading.source_span.start, heading.source_span.end
    ) == (0, 0)
    assert "unclaimed-content=0" in _provenance_message(diagnostics)


def test_source_anchor_aliases_accept_foreign_bookmarks() -> None:
    source = docx_source.read(_CONVENTIONAL)
    ordinals = docx_pandoc.anchored_ordinals(source)
    targets = [f"#pansrc{ordinal}" for ordinal in ordinals]
    targets[0] = "#OLE_LINK7"
    targets[1] = "#_Toc123"

    assert docx_pandoc.source_anchor_aliases(targets, source) == {
        "OLE_LINK7": ordinals[0],
        "_Toc123": ordinals[1],
    }


def test_source_anchor_aliases_reject_one_alias_for_two_ordinals() -> None:
    source = docx_source.read(_CONVENTIONAL)
    ordinals = docx_pandoc.anchored_ordinals(source)
    targets = [f"#pansrc{ordinal}" for ordinal in ordinals]
    targets[0] = targets[1] = "#shared"

    with pytest.raises(ValueError, match="anchor alias 'shared' identifies source ordinals"):
        docx_pandoc.source_anchor_aliases(targets, source)


def test_real_dot_link_paragraph_is_not_mistaken_for_the_farm(tmp_path: Path) -> None:
    """A legitimate paragraph that is exactly one link labelled `.` must survive:
    its injected anchor breaks the farm shape, and even anchor-less it fails the
    pansrc-target requirement."""
    document = Document()
    document.add_paragraph("before")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(".")
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="{ooxml.W_NS}" w:anchor="elsewhere"/>'
    )
    run._r.addprevious(hyperlink)
    hyperlink.append(run._r)
    path = tmp_path / "dotlink.docx"
    document.save(str(path))

    source = docx_source.read(path)
    imported = da.adapt(source, tmp_path / "media", [])

    from pancratius.ir.inlines import inline_plain as plain
    texts = [plain(b.inlines) for b in imported.blocks if isinstance(b, ir.Paragraph) and b.inlines]
    assert texts == ["before", "."]
    # and the pure classifier itself refuses a dot-link paragraph with no pansrc target
    fake = {"t": "Para", "c": [
        {"t": "Link", "c": [["", [], []], [{"t": "Str", "c": "."}], ["https://example.org", ""]]},
    ]}
    assert docx_pandoc.farm_link_targets(fake) is None


def test_phantom_interval_claims_fail_closed(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A span interval covering a content ordinal no anchor proved is known-false
    identity: the import must STOP, not continue on a corrupt claim."""
    document = Document()
    document.add_paragraph("first")
    document.add_paragraph("middle")
    document.add_paragraph("last")
    path = tmp_path / "phantom.docx"
    document.save(str(path))
    source = docx_source.read(path)

    def fused_run_json(
        _src: docx_source.DocxSourceDocument, _media: Path
    ) -> tuple[dict[str, object], str]:
        def anchor(ordinal: int) -> dict[str, object]:
            return {"t": "Span", "c": [[f"pansrc{ordinal}", ["anchor"], []], []]}

        def dot_link(target: str) -> dict[str, object]:
            return {"t": "Link", "c": [["", [], []], [{"t": "Str", "c": "."}], [target, ""]]}

        blocks = [
            # One fused Para claiming ordinals 0 and 2 — ordinal 1's anchor is gone.
            {"t": "Para", "c": [
                {"t": "Str", "c": "first middle last"}, anchor(0), anchor(2),
            ]},
            {"t": "Para", "c": [
                dot_link("#pansrcfarm0"),
                dot_link("#pansrc0"), dot_link("#pansrc1"), dot_link("#pansrc2"),
                {"t": "Span", "c": [["pansrcfarm0", ["anchor"], []], []]},
            ]},
        ]
        return {"blocks": blocks, "meta": {}, "pandoc-api-version": [1, 23, 1]}, ""

    monkeypatch.setattr(docx_pandoc, "run_json", fused_run_json)
    with pytest.raises(da.ProvenanceError, match="claimed by span interval"):
        da.adapt(source, tmp_path / "media", [])


def test_all_heading_document_recovers_every_ordinal(tmp_path: Path) -> None:
    """The hostile farm case: every anchored paragraph is a heading, so Pandoc
    rewrites EVERY ordinal link in the farm. The chunk's reserved marker must
    still identify it — headings keep their spans, no synthetic paragraph leaks."""
    document = Document()
    document.add_heading("First Light", level=1)     # ordinal 0
    document.add_heading("Second Light", level=2)    # ordinal 1
    path = tmp_path / "all-headings.docx"
    document.save(str(path))

    source = docx_source.read(path)
    diagnostics: list[ir.Diagnostic] = []
    imported = da.adapt(source, tmp_path / "media", diagnostics)

    headings = [b for b in imported.blocks if isinstance(b, ir.Heading)]
    assert [
        (span.start, span.end) for heading in headings if (span := heading.source_span)
    ] == [(0, 0), (1, 1)]
    assert not any(isinstance(b, ir.Paragraph) and b.inlines for b in imported.blocks)
    assert "unclaimed-content=0" in _provenance_message(diagnostics)


def test_bookmark_ids_are_unique_across_gaps_and_farm(tmp_path: Path) -> None:
    """Ordinal gaps (empty/pagination paragraphs) must not let a content anchor
    collide with a farm marker: one monotonic allocator covers both."""
    document = Document()
    document.add_paragraph("content zero")   # ordinal 0
    document.add_paragraph()                 # ordinal 1: empty, never anchored
    document.add_paragraph("content two")    # ordinal 2
    path = tmp_path / "id-gaps.docx"
    document.save(str(path))

    source = docx_source.read(path)
    projected = docx_pandoc.project_package(source, tmp_path)
    with zipfile.ZipFile(projected) as archive:
        root = ET.fromstring(archive.read(docx_source.DOCUMENT_PART))
    ids = [
        element.get(f"{docx_source.W}id")
        for element in root.findall(f".//{docx_source.W}bookmarkStart")
    ]
    assert len(ids) == len(set(ids)), f"bookmark ids collide: {sorted(ids)}"
