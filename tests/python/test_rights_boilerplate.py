from __future__ import annotations

import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import parse_xml

from audit.rights_boilerplate import docx_restrictions, markdown_notices
from pancratius import docx_adapter, docx_source, ir, ooxml
from pancratius.docx_optimize import optimize_docx, scrub_document_xml
from pancratius.ir.inlines import inline_plain
from pancratius.passes.scrub import scrub_rights
from pancratius.rights_boilerplate import (
    RightsBoilerplateKind,
    RightsRemovalMismatch,
    classify_rights_boilerplate_notice,
    plan_rights_removal,
)


@pytest.mark.parametrize(
    ("text", "kind"),
    [
        ("© Сергей Орехов (Панкратиус), 2025, 2026", RightsBoilerplateKind.COPYRIGHT),
        ("Copyright © 2025 John Doe", RightsBoilerplateKind.COPYRIGHT),
        ("Copyright © 2025 Pancratius", RightsBoilerplateKind.COPYRIGHT),
        ("All rights reserved.", RightsBoilerplateKind.RESERVED_RIGHTS),
        (
            "No part of this book may be reproduced.",
            RightsBoilerplateKind.REPRODUCTION_RESTRICTION,
        ),
        (
            "No part of this book may be reproduced or utilized without permission.",
            RightsBoilerplateKind.REPRODUCTION_RESTRICTION,
        ),
        (
            "Никакая часть этой книги не может быть воспроизведена без разрешения.",
            RightsBoilerplateKind.REPRODUCTION_RESTRICTION,
        ),
        (
            "The characters and events portrayed are fictitious; any resemblance is coincidental.",
            RightsBoilerplateKind.FICTION_DISCLAIMER,
        ),
        ("Эта книга даруется миру свободно.", RightsBoilerplateKind.FREE_GIFT_NOTICE),
        ("Пусть её чистота будет сохранена.", RightsBoilerplateKind.PURITY_NOTICE),
    ],
)
def test_standalone_notice_classifies_every_kind(
    text: str,
    kind: RightsBoilerplateKind,
) -> None:
    assert classify_rights_boilerplate_notice(text) is kind


@pytest.mark.parametrize(
    "prose",
    [
        'The phrase "all rights reserved" is discussed here.',
        "Copyright © 2025 is discussed here as a historical convention.",
        "© 2025 appears on the archival cover, but this chapter is about provenance.",
        "© 2025 The inscription appears on the archival cover.",
        "© 2025 The Inscription On The Cover",
        "© 2025 This is not a copyright notice but quoted prose.",
        "No part of this book may be reproduced is the clause we reject.",
        (
            "The characters and events portrayed are debated here; "
            "calling them coincidental is misleading."
        ),
        "Никакая часть этой книги описывает, как автор воспроизводит старую дискуссию.",
        "Это не только книга — это передача. Пусть её чистота будет сохранена.",
    ],
)
def test_prose_mentions_are_not_removable_notices(prose: str) -> None:
    assert classify_rights_boilerplate_notice(prose) is None


def test_source_plan_has_one_headmatter_boundary_for_h2_and_long_documents(
    tmp_path: Path,
) -> None:
    h2_document = Document()
    h2_document.add_paragraph("All rights reserved")
    h2_document.add_heading("Subheading", level=2)
    h2_document.add_paragraph("All rights reserved")
    h2_path = tmp_path / "h2.docx"
    h2_document.save(str(h2_path))

    h2_plan = plan_rights_removal(docx_source.read(h2_path))
    assert tuple(removal.ordinal for removal in h2_plan.removals) == (
        docx_source.ParagraphOrdinal(0),
    )

    long_document = Document()
    for ordinal in range(1_000):
        long_document.add_paragraph(
            "All rights reserved" if ordinal == 50 else f"paragraph {ordinal}"
        )
    long_path = tmp_path / "long.docx"
    long_document.save(str(long_path))

    assert not plan_rights_removal(docx_source.read(long_path)).removals


@pytest.mark.parametrize(
    "markdown",
    [
        "All rights reserved.",
        "**All rights reserved.**",
        "> All rights reserved.",
        "> **All rights reserved.**",
        "1. All rights reserved.",
        "1) **All rights reserved.**",
        "- All rights reserved.",
        "* **All rights reserved.**",
        "# All rights reserved.",
        "<strong>All rights reserved.</strong>",
    ],
)
def test_markdown_audit_sees_notice_through_block_presentation(markdown: str) -> None:
    assert markdown_notices(f"---\ntitle: Fixture\n---\n\n{markdown}\n") == (
        (5, "All rights reserved."),
    )


def test_markdown_audit_joins_paragraph_lines_and_ignores_code() -> None:
    markdown = """Ordinary prose.

> No part of this book may be
> reproduced.

```
All rights reserved.
```
"""

    assert markdown_notices(markdown) == (
        (3, "No part of this book may be reproduced."),
    )


def test_markdown_audit_ignores_indented_code_and_reads_setext_heading() -> None:
    markdown = """    All rights reserved.

All rights reserved.
--------------------
"""

    assert markdown_notices(markdown) == ((3, "All rights reserved."),)


@pytest.mark.parametrize(
    ("style", "presentation", "block_type"),
    [
        (None, "plain", ir.Paragraph),
        (None, "strong", ir.Paragraph),
        ("Quote", "plain", ir.QuoteBlock),
        ("List Number", "plain", ir.ListBlock),
        ("List Bullet", "plain", ir.ListBlock),
    ],
)
def test_adapter_removes_one_atomic_styled_notice(
    tmp_path: Path,
    style: str | None,
    presentation: str,
    block_type: type[ir.Block],
) -> None:
    document = Document()
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run("All rights reserved.").bold = presentation == "strong"
    document.add_heading("Chapter", level=1)
    path = tmp_path / "styled.docx"
    document.save(str(path))

    source = docx_source.read(path)
    plan = plan_rights_removal(source)
    adapted = docx_adapter.adapt(source, tmp_path / "media", [])

    assert tuple(removal.ordinal.value for removal in plan.removals) == (0,)
    assert isinstance(adapted.blocks[0], block_type)
    assert adapted.blocks[0].source_span == ir.SourceSpan(0, 0)
    assert scrub_rights(adapted.blocks, plan) == adapted.blocks[1:]


def test_composite_list_notice_fails_closed(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("All rights reserved.", style="List Number")
    document.add_paragraph("All rights reserved.", style="List Number")
    document.add_heading("Chapter", level=1)
    path = tmp_path / "composite-list.docx"
    document.save(str(path))

    source = docx_source.read(path)
    plan = plan_rights_removal(source)
    adapted = docx_adapter.adapt(source, tmp_path / "media", [])

    assert tuple(removal.ordinal.value for removal in plan.removals) == (0, 1)
    assert isinstance(adapted.blocks[0], ir.ListBlock)
    with pytest.raises(RightsRemovalMismatch, match="0:reserved_rights, 1:reserved_rights"):
        scrub_rights(adapted.blocks, plan)


def test_notice_with_opaque_payload_is_not_planned(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph("All rights reserved.")
    bookmark = parse_xml(
        f'<w:bookmarkStart xmlns:w="{ooxml.W_NS}" w:id="1" w:name="keep-me"/>'
    )
    paragraph._p.insert(1, bookmark)
    path = tmp_path / "opaque.docx"
    document.save(str(path))

    source = docx_source.read(path)

    assert source.paragraphs[0].semantics.has_opaque_payload
    assert not plan_rights_removal(source).removals


def test_optimizer_removes_only_selected_standalone_head_notice(tmp_path: Path) -> None:
    document = Document()
    notice = document.add_paragraph()
    notice.add_run("  All rights")
    notice.add_run().add_break(WD_BREAK.PAGE)
    notice.add_run("reserved  ")
    alternate = document.add_paragraph().add_run()
    alternate._r.append(
        parse_xml(
            f'<mc:AlternateContent xmlns:mc="{ooxml.MC_NS}" '
            f'xmlns:w="{ooxml.W_NS}" xmlns:x="urn:unsupported">'
            '<mc:Choice Requires="x"><w:t>All rights reserved</w:t></mc:Choice>'
            "<mc:Fallback><w:t>Visible fallback</w:t></mc:Fallback>"
            "</mc:AlternateContent>"
        )
    )
    document.add_paragraph('The phrase "all rights reserved" is discussed here.')
    document.add_heading("Chapter", level=1)
    document.add_paragraph("All rights reserved")
    path = tmp_path / "rights.docx"
    document.save(str(path))

    source = docx_source.read(path)
    plan = plan_rights_removal(source)
    with zipfile.ZipFile(path) as archive:
        original = archive.read(docx_source.DOCUMENT_PART)
    scrubbed = scrub_document_xml(original, plan)
    root = ET.fromstring(scrubbed)
    texts = [
        docx_source.paragraph_text(paragraph)
        for paragraph in docx_source.story_paragraph_elements(root)
    ]

    assert texts == [
        "Visible fallback",
        'The phrase "all rights reserved" is discussed here.',
        "Chapter",
        "All rights reserved",
    ]
    assert root.find(f".//{ooxml.MC_ALTERNATE_CONTENT}") is not None

    blocks: list[ir.Block] = [
        ir.Paragraph(
            inlines=[ir.Text(paragraph.text)],
            source_span=ir.SourceSpan(paragraph.ordinal.value, paragraph.ordinal.value),
        )
        for paragraph in source.paragraphs
        if paragraph.text
    ]
    cleaned = scrub_rights(blocks, plan)
    assert [
        inline_plain(block.inlines)
        for block in cleaned
        if isinstance(block, ir.Paragraph)
    ] == texts


def test_docx_audit_reads_restrictions_inside_tables(tmp_path: Path) -> None:
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "All rights reserved"
    path = tmp_path / "table.docx"
    document.save(str(path))

    assert docx_restrictions(path) == ((0, "All rights reserved"),)


def test_optimizer_scrubs_without_document_relationship_part(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("All rights reserved")
    document.add_heading("Chapter", level=1)
    original = tmp_path / "original.docx"
    source = tmp_path / "without-rels.docx"
    output = tmp_path / "optimized.docx"
    document.save(str(original))

    with zipfile.ZipFile(original) as archive, zipfile.ZipFile(
        source, "w", zipfile.ZIP_DEFLATED
    ) as rewritten:
        for info in archive.infolist():
            if info.filename != "word/_rels/document.xml.rels":
                rewritten.writestr(info, archive.read(info.filename))

    optimize_docx(source, output)

    assert [paragraph.text for paragraph in docx_source.read(output).paragraphs] == [
        "Chapter"
    ]
