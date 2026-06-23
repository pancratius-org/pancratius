from __future__ import annotations

import base64
import re
import shutil
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import pytest

from pancratius.ooxml import NamespaceBinding, serialize_xml
from pancratius.translation.docx import (
    DocxTranslationError,
    print_batch,
    render_translated_docx,
    translate_docx_batch,
)
from pancratius.translation.docx.pipeline import (
    _dedupe_media_payloads,
    _repair_unbound_relationship_prefixes,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


requires_pandoc = pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc required")


def _write_source_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("ГЛАВА 1", level=1)
    doc.add_paragraph("Панкратиус: Привет")
    doc.add_paragraph("Первая строка")
    doc.add_paragraph("Вторая строка")
    doc.add_paragraph("")
    doc.add_paragraph("Финал")
    doc.save(str(path))


def _write_paragraph_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(str(path))


def _write_superscript_first_run_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    paragraph = doc.add_paragraph()
    first = paragraph.add_run("С")
    first.font.superscript = True
    paragraph.add_run("вет")
    doc.save(str(path))


def _write_superscript_character_style_first_run_docx(path: Path) -> None:
    _write_paragraph_docx(path, ["Свет"])
    with zipfile.ZipFile(path) as zf:
        parts = {name: zf.read(name) for name in zf.namelist()}
    styles = ET.fromstring(parts["word/styles.xml"])
    style = ET.SubElement(styles, f"{W}style")
    style.set(f"{W}type", "character")
    style.set(f"{W}styleId", "af2")
    name = ET.SubElement(style, f"{W}name")
    name.set(f"{W}val", "footnote reference")
    rpr = ET.SubElement(style, f"{W}rPr")
    vert_align = ET.SubElement(rpr, f"{W}vertAlign")
    vert_align.set(f"{W}val", "superscript")

    root = ET.fromstring(parts["word/document.xml"])
    run = root.find(f".//{W}body/{W}p/{W}r")
    assert run is not None
    run_rpr = run.find(f"{W}rPr")
    if run_rpr is None:
        run_rpr = ET.Element(f"{W}rPr")
        run.insert(0, run_rpr)
    rstyle = ET.SubElement(run_rpr, f"{W}rStyle")
    rstyle.set(f"{W}val", "af2")

    parts["word/styles.xml"] = ET.tostring(styles, encoding="UTF-8", xml_declaration=True)
    parts["word/document.xml"] = ET.tostring(root, encoding="UTF-8", xml_declaration=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in parts.items():
            zf.writestr(name, data)


def _write_field_cover_docx(path: Path, old_data_uri: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Начало")
    doc.add_paragraph("")
    doc.add_paragraph("Финал")
    doc.save(str(path))

    with zipfile.ZipFile(path) as zf:
        parts = {name: zf.read(name) for name in zf.namelist()}
    root = ET.fromstring(parts["word/document.xml"])
    body = root.find(f"{W}body")
    assert body is not None
    cover_paragraph = body.findall(f"{W}p")[1]
    run = ET.SubElement(cover_paragraph, f"{W}r")
    instr = ET.SubElement(run, f"{W}instrText")
    instr.text = f' INCLUDEPICTURE "{old_data_uri}" '
    parts["word/document.xml"] = ET.tostring(root, encoding="UTF-8", xml_declaration=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in parts.items():
            zf.writestr(name, data)


def _write_footnote_anchor_docx(path: Path) -> None:
    _write_paragraph_docx(path, ["Свет"])
    with zipfile.ZipFile(path) as zf:
        parts = {name: zf.read(name) for name in zf.namelist()}
    root = ET.fromstring(parts["word/document.xml"])
    paragraph = root.find(f".//{W}body/{W}p")
    assert paragraph is not None
    run = ET.SubElement(paragraph, f"{W}r")
    ref = ET.SubElement(run, f"{W}footnoteReference")
    ref.set(f"{W}id", "1")
    parts["word/document.xml"] = ET.tostring(root, encoding="UTF-8", xml_declaration=True)
    parts["word/footnotes.xml"] = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<w:footnotes xmlns:w="{W_NS}">'
        f'<w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r>'
        f"<w:r><w:t>Сноска</w:t></w:r></w:p></w:footnote>"
        f"</w:footnotes>"
    ).encode()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in parts.items():
            zf.writestr(name, data)


def _write_two_footnote_anchor_docx(path: Path) -> None:
    _write_paragraph_docx(path, ["и тьма не объяла его", "Ин. 1:5", "Не бойся прикасаться."])
    with zipfile.ZipFile(path) as zf:
        parts = {name: zf.read(name) for name in zf.namelist()}
    root = ET.fromstring(parts["word/document.xml"])
    paragraphs = root.findall(f".//{W}body/{W}p")
    for index, paragraph in enumerate((paragraphs[0], paragraphs[2]), start=1):
        run = ET.SubElement(paragraph, f"{W}r")
        ref = ET.SubElement(run, f"{W}footnoteReference")
        ref.set(f"{W}id", str(index))
    parts["word/document.xml"] = ET.tostring(root, encoding="UTF-8", xml_declaration=True)
    parts["word/footnotes.xml"] = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<w:footnotes xmlns:w="{W_NS}">'
        f'<w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r>'
        f"<w:r><w:t>Первая сноска</w:t></w:r></w:p></w:footnote>"
        f'<w:footnote w:id="2"><w:p><w:r><w:footnoteRef/></w:r>'
        f"<w:r><w:t>Вторая сноска</w:t></w:r></w:p></w:footnote>"
        f"</w:footnotes>"
    ).encode()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in parts.items():
            zf.writestr(name, data)


def _write_image_text_docx(path: Path, image_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    image_path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    ))
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(image_path), width=Inches(0.1))
    paragraph.add_run("Текст")
    doc.save(str(path))


def _write_source_only_image_docx(path: Path, image_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    image_path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    ))
    doc = Document()
    doc.add_paragraph("Начало")
    doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.1))
    doc.add_paragraph("Финал")
    doc.save(str(path))


def _write_tiny_png(path: Path) -> None:
    path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    ))


def _paragraph_text_and_style(path: Path) -> list[tuple[str, str]]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    rows: list[tuple[str, str]] = []
    for p in root.findall(f".//{W}body/{W}p"):
        text = "".join(t.text or "" for t in p.findall(f".//{W}t"))
        style = p.find(f"./{W}pPr/{W}pStyle")
        rows.append((text, "" if style is None else str(style.get(f"{W}val") or "")))
    return rows


def _paragraph_vert_align_values(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    out: list[str] = []
    for vert_align in root.findall(f".//{W}body/{W}p/{W}r/{W}rPr/{W}vertAlign"):
        out.append(str(vert_align.get(f"{W}val") or ""))
    return out


def _paragraph_run_style_values(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    out: list[str] = []
    for style in root.findall(f".//{W}body/{W}p/{W}r/{W}rPr/{W}rStyle"):
        out.append(str(style.get(f"{W}val") or ""))
    return out


def _write_md(path: Path, body: str) -> None:
    path.write_text(body, encoding="utf-8")


def _write_catalog_book_md(path: Path, *, number: int, lang: str, title: str, body: str) -> None:
    _write_md(
        path,
        f"""---
kind: book
number: {number}
slug: test
title: {title}
lang: {lang}
description: Test
---

{body}
""",
    )


def _assert_document_prefix_safety(docx: Path) -> None:
    with zipfile.ZipFile(docx) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        infos = zf.infolist()

    assert all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in infos)
    assert re.search(r"<w:document\b", xml)
    assert not re.search(r"\bxmlns:ns\d+=", xml)
    assert not re.search(r"</?ns\d+:", xml)
    ignorable = re.search(r'\bmc:Ignorable="([^"]+)"', xml)
    assert ignorable is not None
    for prefix in ignorable.group(1).split():
        assert f"xmlns:{prefix}=" in xml


def test_ooxml_serialize_xml_does_not_leak_elementtree_namespace_registry() -> None:
    namespace_map = getattr(ET, "_namespace_map", None)
    assert isinstance(namespace_map, dict)
    before = dict(namespace_map)
    root = ET.Element("{http://example.com/pancratius-test}root")

    payload = serialize_xml(
        root,
        bindings=(NamespaceBinding("pan", "http://example.com/pancratius-test"),),
    )

    assert b"<pan:root" in payload
    assert dict(namespace_map) == before


@requires_pandoc
def test_render_translated_docx_preserves_word_slots_and_transfers_english(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_source_docx(source_docx)
    _write_md(
        source_md,
        """## ГЛАВА 1

**Панкратиус:**

Привет

<div class="lineated verse">

Первая строка  
Вторая строка

Финал

</div>
""",
    )
    _write_md(
        translated_md,
        """## CHAPTER 1

**Pankratius:**

Hello

<div class="lineated verse">

First line  
Second line

Final line

</div>
""",
    )

    source_units, translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert (source_units, translated_units, aligned_units) == (7, 7, 7)
    assert _paragraph_text_and_style(out) == [
        ("CHAPTER 1", "Heading1"),
        ("Pankratius: Hello", ""),
        ("First line", ""),
        ("Second line", ""),
        ("", ""),
        ("Final line", ""),
    ]


@requires_pandoc
def test_render_translated_docx_is_deterministic_and_preserves_ooxml_prefixes(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out_a = tmp_path / "a.docx"
    out_b = tmp_path / "b.docx"
    _write_paragraph_docx(source_docx, ["Свет"])
    _write_md(source_md, "Свет\n")
    _write_md(translated_md, "Light\n")

    first = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out_a,
    )
    second = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out_b,
    )

    assert [d for d in first[3] if d.severity == "fatal"] == []
    assert [d for d in second[3] if d.severity == "fatal"] == []
    assert out_a.read_bytes() == out_b.read_bytes()
    _assert_document_prefix_safety(out_a)


@requires_pandoc
def test_render_translated_docx_refuses_skipped_source_text_slot(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Первый", "Лишний абзац", "Второй"])
    _write_md(source_md, "Первый\n\nВторой\n")
    _write_md(translated_md, "First\n\nSecond\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert aligned_units == 0
    assert any(
        diagnostic.severity == "fatal"
        and diagnostic.code == "docx-translate.transfer-failed"
        and "Лишний абзац" in diagnostic.message
        for diagnostic in diagnostics
    )
    assert not out.exists()


@requires_pandoc
def test_render_translated_docx_removes_source_only_toc(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Оглавление", "Глава 1 1", "Начало"])
    _write_md(source_md, "Начало\n")
    _write_md(translated_md, "Beginning\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert any(
        diagnostic.severity == "warning"
        and diagnostic.code == "docx-translate.source-slot-removed"
        and "source-toc" in diagnostic.message
        for diagnostic in diagnostics
    )
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Beginning", "")]


@requires_pandoc
def test_render_translated_docx_removes_trailing_source_only_back_matter(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Финал", "Копирайт", "© Сергей Орехов"])
    _write_md(source_md, "Финал\n")
    _write_md(translated_md, "Final\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert any(
        diagnostic.severity == "warning"
        and diagnostic.code == "docx-translate.source-slot-removed"
        and "source-back-matter" in diagnostic.message
        for diagnostic in diagnostics
    )
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Final", "")]


@requires_pandoc
def test_render_translated_docx_aligns_markdown_superscript_markup(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["∣R∣=2ℵ0>ℵ0"])
    _write_md(source_md, "∣R∣=2^ℵ0^>ℵ~0~\n")
    _write_md(translated_md, "∣R∣=2^ℵ0^>ℵ~0~\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("∣R∣=2ℵ0>ℵ0", "")]


@requires_pandoc
def test_render_translated_docx_does_not_inherit_source_superscript_for_body_text(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_superscript_first_run_docx(source_docx)
    _write_md(source_md, "Свет\n")
    _write_md(translated_md, "Light\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Light", "")]
    assert _paragraph_vert_align_values(out) == []


@requires_pandoc
def test_render_translated_docx_does_not_inherit_source_character_style_for_body_text(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_superscript_character_style_first_run_docx(source_docx)
    _write_md(source_md, "Свет\n")
    _write_md(translated_md, "Light\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Light", "")]
    assert _paragraph_run_style_values(out) == []


@requires_pandoc
def test_render_translated_docx_aligns_smiley_missing_from_source_docx(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Панкратиус: Нет моего шага"])
    _write_md(source_md, "Панкратиус: ☺ Нет моего шага\n")
    _write_md(translated_md, "Pankratius: ☺ There is no my step\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Pankratius: ☺ There is no my step", "")]


@requires_pandoc
def test_render_translated_docx_preserves_hyperlinks(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["ссылка"])
    _write_md(source_md, "[ссылка](https://example.com)\n")
    _write_md(translated_md, "[link](https://example.org)\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    with zipfile.ZipFile(out) as zf:
        document = ET.fromstring(zf.read("word/document.xml"))
        rels = ET.fromstring(zf.read("word/_rels/document.xml.rels"))
    hyperlink = document.find(f".//{W}hyperlink")
    assert hyperlink is not None
    rel_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
    assert rel_id
    rel = rels.find(
        "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
        f"[@Id='{rel_id}']"
    )
    assert rel is not None
    assert rel.get("Target") == "https://example.org"
    assert rel.get("TargetMode") == "External"


@requires_pandoc
def test_render_translated_docx_refuses_raw_inline_html(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Текст"])
    _write_md(source_md, "<u>Текст</u>\n")
    _write_md(translated_md, "<u>Text</u>\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert aligned_units == 0
    assert any(
        diagnostic.severity == "fatal"
        and diagnostic.code == "docx-translate.raw-inline-html-skipped"
        for diagnostic in diagnostics
    )
    assert not out.exists()


@requires_pandoc
def test_render_translated_docx_allows_empty_scripture_wrappers(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Истина"])
    _write_md(source_md, '<blockquote class="scripture">\n\nИстина\n\n</blockquote>\n')
    _write_md(translated_md, '<blockquote class="scripture">\n\nTruth\n\n</blockquote>\n')

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity in {"fatal", "warning"}] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Truth", "")]


@requires_pandoc
def test_render_translated_docx_splits_blockquote_lines_into_word_paragraphs(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Первая строка", "Вторая строка"])
    _write_md(source_md, "> *Первая строка*\n> *Вторая строка*\n")
    _write_md(translated_md, "> *First line*\n> *Second line*\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("First line", ""), ("Second line", "")]


@requires_pandoc
def test_render_translated_docx_splits_lineated_soft_breaks(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Первая строка", "Вторая строка"])
    _write_md(source_md, '<div class="lineated">\n\nПервая строка\nВторая строка\n\n</div>\n')
    _write_md(translated_md, '<div class="lineated">\n\nFirst line\nSecond line\n\n</div>\n')

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("First line", ""), ("Second line", "")]


@requires_pandoc
def test_render_translated_docx_treats_braille_blank_as_blank_line(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, [""])
    _write_md(source_md, '<div class="lineated">\n\n⠀\n\n</div>\n')
    _write_md(translated_md, '<div class="lineated">\n\n⠀\n\n</div>\n')

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("", "")]


@requires_pandoc
def test_render_translated_docx_joins_dialogue_label_with_colon_continuation(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Панкратиус: Отец, продолжим?"])
    _write_md(source_md, "**Панкратиус:**\n\n: Отец, продолжим?\n")
    _write_md(translated_md, "**Pankratius:**\n\n: Father, shall we continue?\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Pankratius: Father, shall we continue?", "")]


@requires_pandoc
def test_render_translated_docx_aligns_label_when_source_docx_omits_terminal_colon(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Светозар"])
    _write_md(source_md, "**Светозар:**\n")
    _write_md(translated_md, "**Svetozar:**\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Svetozar:", "")]


@requires_pandoc
def test_render_translated_docx_aligns_dialogue_colon_before_dash(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Светозар — это имя Света"])
    _write_md(source_md, "**Светозар:**\n\n— это имя Света\n")
    _write_md(translated_md, "**Svetozar:**\n\n— this is the name of Light\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Svetozar: — this is the name of Light", "")]


@requires_pandoc
def test_render_translated_docx_aligns_colon_emoticon_continuation(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Панкратиус::)))) Продолжай"])
    _write_md(source_md, "**Панкратиус:**\n\n:)))) Продолжай\n")
    _write_md(translated_md, "**Pankratius:**\n\n:)))) Continue\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Pankratius::)))) Continue", "")]


@requires_pandoc
def test_render_translated_docx_aligns_known_split_letter_typo(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Панкратиус: Пока хочу оформить"])
    _write_md(source_md, "**Панкратиус:**\n\nП ока хочу оформить\n")
    _write_md(translated_md, "**Pankratius:**\n\nFor now I want to prepare it\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Pankratius: For now I want to prepare it", "")]


@requires_pandoc
@pytest.mark.parametrize("word", ["так", "как", "храм"])
def test_render_translated_docx_aligns_spurious_hyphen_before_russian_connector(
    tmp_path: Path,
    word: str,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, [f"третья часть дня не светла была {word}, как и ночи."])
    _write_md(source_md, f"третья часть дня не светла была ‑{word}, как и ночи.\n")
    _write_md(translated_md, "a third of the day was not bright, and likewise the night.\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [
        ("a third of the day was not bright, and likewise the night.", "")
    ]


@requires_pandoc
def test_render_translated_docx_aligns_one_spurious_connector_hyphen_without_dropping_real_dash(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["ноги у него как у медведя, а пасть у него - как пасть у льва."])
    _write_md(source_md, "ноги у него ‑как у медведя, а пасть у него - как пасть у льва.\n")
    _write_md(translated_md, "its feet were like a bear's, and its mouth like a lion's mouth.\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [
        ("its feet were like a bear's, and its mouth like a lion's mouth.", "")
    ]


@requires_pandoc
def test_render_translated_docx_aligns_nonbreaking_hyphen_dropped_by_docx(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Не ктото смотрел."])
    _write_md(source_md, "Не кто‑то смотрел.\n")
    _write_md(translated_md, "It was not someone who looked.\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("It was not someone who looked.", "")]


@requires_pandoc
def test_render_translated_docx_aligns_source_docx_scraped_citation_suffix(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Калки восстановит дхарму. Википедия+2Википедия+2"])
    _write_md(source_md, "Калки восстановит дхарму.\n")
    _write_md(translated_md, "Kalki will restore dharma.\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 1
    assert _paragraph_text_and_style(out) == [("Kalki will restore dharma.", "")]


@requires_pandoc
def test_render_translated_docx_refuses_unproven_intra_word_hyphen_equivalence(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Godman"])
    _write_md(source_md, "God-man\n")
    _write_md(translated_md, "God-man\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert aligned_units == 0
    assert any(
        diagnostic.severity == "fatal"
        and diagnostic.code == "docx-translate.transfer-failed"
        for diagnostic in diagnostics
    )
    assert not out.exists()


@requires_pandoc
def test_render_translated_docx_aligns_thematic_break_to_blank_word_paragraph(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Начало", "", "Финал"])
    _write_md(source_md, "Начало\n\n***\n\nФинал\n")
    _write_md(translated_md, "Beginning\n\n***\n\nFinal\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 3
    assert _paragraph_text_and_style(out) == [("Beginning", ""), ("***", ""), ("Final", "")]


@requires_pandoc
def test_render_translated_docx_preserves_literal_thematic_break_from_word(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Начало", "* * *", "Финал"])
    _write_md(source_md, "Начало\n\n***\n\nФинал\n")
    _write_md(translated_md, "Beginning\n\n***\n\nFinal\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 3
    assert _paragraph_text_and_style(out) == [("Beginning", ""), ("* * *", ""), ("Final", "")]


@requires_pandoc
def test_render_translated_docx_removes_source_only_thematic_separator_with_warning(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Начало", "***", "Финал"])
    _write_md(source_md, "Начало\n\nФинал\n")
    _write_md(translated_md, "Beginning\n\nFinal\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert any(
        diagnostic.severity == "warning"
        and diagnostic.code == "docx-translate.source-slot-removed"
        and "source-thematic-separator" in diagnostic.message
        for diagnostic in diagnostics
    )
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Beginning", ""), ("Final", "")]


@requires_pandoc
def test_render_translated_docx_skips_markdown_blank_absent_from_source_docx(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Проверь себя:", "1. Да"])
    _write_md(source_md, '<div class="lineated">\n\nПроверь себя:\n\n1\\. Да\n\n</div>\n')
    _write_md(translated_md, '<div class="lineated">\n\nTest yourself:\n\n1\\. Yes\n\n</div>\n')

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Test yourself:", ""), ("1. Yes", "")]


@requires_pandoc
def test_render_translated_docx_splits_signature_html_lines(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_paragraph_docx(source_docx, ["Сергей Панкратиус", "— через Светозара"])
    _write_md(source_md, """<p class="signature">
Сергей Панкратиус
— через Светозара
</p>
""")
    _write_md(translated_md, """<p class="signature">
Sergei Pankratius
— through Svetozar
</p>
""")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    assert _paragraph_text_and_style(out) == [("Sergei Pankratius", ""), ("— through Svetozar", "")]


@requires_pandoc
def test_render_translated_docx_wraps_invalid_docx_as_diagnostic(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    source_docx.write_bytes(b"not a zip file")
    _write_md(source_md, "Текст\n")
    _write_md(translated_md, "Text\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert aligned_units == 0
    assert any(
        diagnostic.severity == "fatal"
        and diagnostic.code == "docx-translate.invalid-docx"
        for diagnostic in diagnostics
    )
    assert not out.exists()


@requires_pandoc
def test_render_translated_docx_preserves_footnote_body_marker(tmp_path: Path) -> None:
    book_dir = Path("src/content/books/16-velikaya-piramida-i-evangelie-tsarstviya")
    if not book_dir.is_dir():
        pytest.skip("book 16 corpus fixture is required")
    out = tmp_path / "en.docx"

    _source_units, _translated_units, _aligned_units, diagnostics = render_translated_docx(
        source_docx=book_dir / "ru.docx",
        source_md=book_dir / "ru.md",
        translated_md=book_dir / "en.md",
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/footnotes.xml"))
    body_notes = [
        note for note in root.findall(f"{W}footnote")
        if int(note.get(f"{W}id", "0")) > 0
    ]
    assert body_notes
    assert body_notes[0].find(f".//{W}footnoteRef") is not None


@requires_pandoc
def test_render_translated_docx_reads_footnote_anchors_inside_raw_epigraph_html(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_two_footnote_anchor_docx(source_docx)
    _write_md(
        source_md,
        """<blockquote class="epigraph">
<p>
и тьма не объяла его[^1]
</p>
<footer>
Ин. 1:5
</footer>
</blockquote>

Не бойся прикасаться[^2].

[^1]: Первая сноска

[^2]: Вторая сноска
""",
    )
    _write_md(
        translated_md,
        """<blockquote class="epigraph">
<p>
and the darkness has not overcome it[^1]
</p>
<footer>
Jn. 1:5
</footer>
</blockquote>

Do not be afraid to touch[^2].

[^1]: First footnote

[^2]: Second footnote
""",
    )

    source_units, translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert (source_units, translated_units, aligned_units) == (3, 3, 3)
    assert _paragraph_text_and_style(out) == [
        ("and the darkness has not overcome it", ""),
        ("Jn. 1:5", ""),
        ("Do not be afraid to touch.", ""),
    ]
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/footnotes.xml"))
    body_notes = [
        "".join(t.text or "" for t in note.findall(f".//{W}t"))
        for note in root.findall(f"{W}footnote")
        if int(note.get(f"{W}id", "0")) > 0
    ]
    assert body_notes == ["First footnote", "Second footnote"]


@requires_pandoc
def test_render_translated_docx_refuses_source_anchor_missing_from_markdown(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_footnote_anchor_docx(source_docx)
    _write_md(source_md, "Свет\n\n[^1]: Сноска\n")
    _write_md(translated_md, "Light\n\n[^1]: Footnote\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert aligned_units == 0
    assert any(
        diagnostic.severity == "fatal"
        and diagnostic.code == "docx-translate.transfer-failed"
        and "source has 1" in diagnostic.message
        for diagnostic in diagnostics
    )
    assert not out.exists()


@requires_pandoc
def test_render_translated_docx_transfers_image_alt_metadata(tmp_path: Path) -> None:
    book_dir = Path("src/content/books/75-otkrovenie-tvortsa-dlya-ii")
    if not book_dir.is_dir():
        pytest.skip("book 75 corpus fixture is required")
    out = tmp_path / "en.docx"

    _source_units, _translated_units, _aligned_units, diagnostics = render_translated_docx(
        source_docx=book_dir / "ru.docx",
        source_md=book_dir / "ru.md",
        translated_md=book_dir / "en.md",
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    metadata_nodes = [
        node for node in root.iter() if node.tag.rsplit("}", 1)[-1] in {"docPr", "cNvPr"}
    ]
    assert metadata_nodes
    assert metadata_nodes[0].get("name") == "Illustration"
    assert metadata_nodes[0].get("descr") == "Illustration"
    assert any(node.tag.rsplit("}", 1)[-1] == "cNvPr" for node in metadata_nodes)
    assert not re.search(r"[А-Яа-яЁё]", "".join(str(node.attrib) for node in metadata_nodes))


@requires_pandoc
def test_render_translated_docx_preserves_image_when_word_slot_also_has_text(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_image_text_docx(source_docx, tmp_path / "pixel.png")
    _write_md(source_md, "![Иллюстрация](./pixel.png)\n\nТекст\n")
    _write_md(translated_md, "![Illustration](./pixel.png)\n\nText\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    assert root.find(f".//{W}drawing") is not None
    assert _paragraph_text_and_style(out) == [("Text", "")]


@requires_pandoc
def test_render_translated_docx_preserves_source_only_drawing_without_matching_blank(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_source_only_image_docx(source_docx, tmp_path / "pixel.png")
    _write_md(source_md, "Начало\n\nФинал\n")
    _write_md(translated_md, "Beginning\n\nFinal\n")

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    assert root.find(f".//{W}drawing") is not None
    assert _paragraph_text_and_style(out) == [("Beginning", ""), ("", ""), ("Final", "")]


@requires_pandoc
def test_render_translated_docx_skips_lineated_blank_before_source_only_drawing(
    tmp_path: Path,
) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    out = tmp_path / "en.docx"
    _write_source_only_image_docx(source_docx, tmp_path / "pixel.png")
    _write_md(source_md, '<div class="lineated">\n\nНачало\n\nФинал\n\n</div>\n')
    _write_md(translated_md, '<div class="lineated">\n\nBeginning\n\nFinal\n\n</div>\n')

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert aligned_units == 2
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    assert root.find(f".//{W}drawing") is not None
    assert _paragraph_text_and_style(out) == [("Beginning", ""), ("", ""), ("Final", "")]


@requires_pandoc
def test_render_translated_docx_sanitizes_unmatched_drawing_metadata(tmp_path: Path) -> None:
    book_dir = Path("src/content/books/12-kniga-kotoruyu-nikto-ne-napisal")
    if not book_dir.is_dir():
        pytest.skip("book 12 corpus fixture is required")
    out = tmp_path / "en.docx"

    _source_units, _translated_units, _aligned_units, diagnostics = render_translated_docx(
        source_docx=book_dir / "ru.docx",
        source_md=book_dir / "ru.md",
        translated_md=book_dir / "en.md",
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    with zipfile.ZipFile(out) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    metadata_text = "".join(
        str(node.attrib)
        for node in root.iter()
        if node.tag.rsplit("}", 1)[-1] in {"docPr", "cNvPr"}
    )
    assert not re.search(r"[А-Яа-яЁё]", metadata_text)


@requires_pandoc
def test_render_translated_docx_replaces_embedded_source_cover_data_uri(tmp_path: Path) -> None:
    source_docx = tmp_path / "ru.docx"
    source_md = tmp_path / "ru.md"
    translated_md = tmp_path / "en.md"
    cover = tmp_path / "cover.en.jpg"
    out = tmp_path / "en.docx"
    old_data_uri = "data:image/png;base64,b2xk"
    cover.write_bytes(b"\xff\xd8\xff\xe0new-cover")
    _write_field_cover_docx(source_docx, old_data_uri)
    _write_md(
        source_md,
        """---
cover: ./cover.ru.jpg
---

Начало

Финал
""",
    )
    _write_md(
        translated_md,
        """---
cover: ./cover.en.jpg
---

Beginning

Final
""",
    )

    _source_units, _translated_units, aligned_units, diagnostics = render_translated_docx(
        source_docx=source_docx,
        source_md=source_md,
        translated_md=translated_md,
        out=out,
    )

    assert [d for d in diagnostics if d.severity == "fatal"] == []
    assert any(
        diagnostic.severity == "warning"
        and diagnostic.code == "docx-translate.cover-image-replaced"
        for diagnostic in diagnostics
    )
    assert aligned_units == 2
    with zipfile.ZipFile(out) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert old_data_uri not in xml
    assert "data:image/jpeg;base64," in xml


def test_translate_docx_batch_errors_for_explicit_book_without_prerequisites(tmp_path: Path) -> None:
    content_root = tmp_path / "content"
    book_dir = content_root / "books" / "01-test"
    book_dir.mkdir(parents=True)
    _write_md(
        book_dir / "ru.md",
        """---
kind: book
number: 1
slug: test
title: Test
lang: ru
description: Test
---

Текст
""",
    )

    with pytest.raises(DocxTranslationError, match="book-01 cannot be translated"):
        translate_docx_batch(content_root=content_root, book=1, lang="en", dry_run=True)


def test_translate_docx_batch_refuses_source_locale(tmp_path: Path) -> None:
    with pytest.raises(DocxTranslationError, match="refuses source locale"):
        translate_docx_batch(content_root=tmp_path, lang="ru", dry_run=True)


@requires_pandoc
def test_translate_docx_batch_treats_existing_translated_docx_as_source(
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    content_root = tmp_path / "content"
    book_dir = content_root / "books" / "01-test"
    book_dir.mkdir(parents=True)
    _write_paragraph_docx(book_dir / "ru.docx", ["Свет"])
    _write_catalog_book_md(book_dir / "ru.md", number=1, lang="ru", title="Тест", body="Свет\n")
    _write_catalog_book_md(book_dir / "en.md", number=1, lang="en", title="Test", body="Light\n")

    dry = translate_docx_batch(content_root=content_root, lang="en", dry_run=True)

    assert len(dry.reports) == 1
    assert dry.discovery.missing == 1
    assert dry.discovery.existing == 0
    assert not (book_dir / "en.docx").exists()

    written = translate_docx_batch(content_root=content_root, lang="en")
    assert not written.failed
    assert (book_dir / "en.docx").is_file()

    after = translate_docx_batch(content_root=content_root, lang="en", dry_run=True)
    assert after.reports == ()
    assert after.discovery.missing == 0
    assert after.discovery.existing == 1

    explicit = translate_docx_batch(content_root=content_root, book=1, lang="en", dry_run=True)
    assert explicit.failed
    assert explicit.reports[0].write_report.diagnostics[0].code == "docx-translate.overwrite-refused"
    print_batch(explicit, dry_run=True)
    stdout = capsys.readouterr().out
    assert "REFUSE book-01: would refuse" in stdout
    with pytest.raises(DocxTranslationError, match="--replace requires an explicit book:NN"):
        translate_docx_batch(content_root=content_root, lang="en", dry_run=True, replace=True)
    with pytest.raises(DocxTranslationError, match="--limit must be non-negative"):
        translate_docx_batch(content_root=content_root, lang="en", dry_run=True, limit=-1)
    with pytest.raises(DocxTranslationError, match="--limit cannot be combined"):
        translate_docx_batch(content_root=content_root, book=1, lang="en", dry_run=True, limit=1)


def test_markdown_render_backend_repairs_unbound_relationship_prefix() -> None:
    xml = (
        b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        b'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        b'<w:footerReference ns11:id="rId1" w:type="default" />'
        b"</w:document>"
    )

    repaired = _repair_unbound_relationship_prefixes(xml)

    assert b' r:id="rId1"' in repaired
    assert b"ns11:id" not in repaired


def test_markdown_render_backend_dedupes_media_payloads_and_retargets_relationships() -> None:
    parts = {
        "word/media/rId1.png": b"same",
        "word/media/image1.png": b"same",
        "word/media/image2.png": b"different",
        "word/_rels/document.xml.rels": (
            b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            b'<Relationship Id="rId1" Type="image" Target="media/rId1.png" />'
            b'<Relationship Id="rId2" Type="image" Target="media/image2.png" />'
            b"</Relationships>"
        ),
    }

    _dedupe_media_payloads(parts)

    assert "word/media/image1.png" in parts
    assert "word/media/rId1.png" not in parts
    rels = parts["word/_rels/document.xml.rels"]
    assert b'Target="media/image1.png"' in rels
    assert b"media/rId1.png" not in rels


@requires_pandoc
def test_translate_docx_batch_markdown_render_backend_writes_docx(
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    content_root = tmp_path / "content"
    book_dir = content_root / "books" / "01-test"
    book_dir.mkdir(parents=True)
    _write_paragraph_docx(book_dir / "ru.docx", ["Свет"])
    _write_catalog_book_md(book_dir / "ru.md", number=1, lang="ru", title="Тест", body="Свет\n")
    _write_catalog_book_md(book_dir / "en.md", number=1, lang="en", title="Test", body="Light…\n")

    batch = translate_docx_batch(
        content_root=content_root,
        book=1,
        lang="en",
        backend="markdown-render",
    )

    assert not batch.failed
    assert batch.reports[0].backend == "markdown-render"
    assert (book_dir / "en.docx").is_file()
    with zipfile.ZipFile(book_dir / "en.docx") as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        infos = zf.infolist()
    assert all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in infos)
    assert re.search(r"<w:document\b", xml)
    assert not re.search(r"\bxmlns:ns\d+=", xml)
    assert not re.search(r"</?ns\d+:", xml)
    print_batch(batch, dry_run=False)
    stdout = capsys.readouterr().out
    assert "markdown-render" in stdout


@requires_pandoc
def test_markdown_render_backend_keeps_standalone_image_alt_out_of_body_text(tmp_path: Path) -> None:
    content_root = tmp_path / "content"
    book_dir = content_root / "books" / "01-test"
    images_dir = book_dir / "images"
    images_dir.mkdir(parents=True)
    _write_paragraph_docx(book_dir / "ru.docx", ["Свет"])
    _write_tiny_png(images_dir / "pixel.png")
    _write_catalog_book_md(book_dir / "ru.md", number=1, lang="ru", title="Тест", body="Свет\n")
    _write_catalog_book_md(
        book_dir / "en.md",
        number=1,
        lang="en",
        title="Test",
        body="Before\n\n![Corpus diagram](./images/pixel.png)\n\nAfter\n",
    )

    batch = translate_docx_batch(
        content_root=content_root,
        book=1,
        lang="en",
        backend="markdown-render",
    )

    assert not batch.failed
    with zipfile.ZipFile(book_dir / "en.docx") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
        media_parts = [name for name in zf.namelist() if name.startswith("word/media/")]
    body_text = "\n".join(
        "".join(t.text or "" for t in paragraph.findall(f".//{W}t"))
        for paragraph in root.findall(f".//{W}body/{W}p")
    )
    metadata_text = " ".join(
        str(value)
        for node in root.iter()
        if node.tag.rsplit("}", 1)[-1] in {"docPr", "cNvPr"}
        for value in node.attrib.values()
    )
    assert "Before" in body_text
    assert "After" in body_text
    assert "Corpus diagram" not in body_text
    assert "Corpus diagram" in metadata_text
    assert media_parts
