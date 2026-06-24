from __future__ import annotations

import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import pytest

from pancratius import cli
from pancratius.docx_merge import (
    DocxMergeError,
    merge_docx,
    validate_docx_package,
)
from pancratius.docx_outline import PartBoundary
from pancratius.ooxml import R_NS, REL_NS

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
OFFICE_DOCUMENT_REL = f"{R_NS}/officeDocument"


def _write_png(path: Path, rgb: tuple[int, int, int]) -> None:
    from PIL import Image

    image = Image.new("RGB", (4, 4), rgb)
    image.save(path)


def _write_docx(path: Path, blocks: list[tuple[str, str]], image: Path | None = None) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    for kind, text in blocks:
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "p":
            doc.add_paragraph(text)
        else:  # pragma: no cover - guards test fixture calls
            raise AssertionError(f"unknown block kind: {kind}")
    if image is not None:
        doc.add_picture(str(image), width=Inches(0.25))
    doc.save(str(path))


def _paragraphs(path: Path) -> list[tuple[str, str]]:
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    rows: list[tuple[str, str]] = []
    for p in root.findall(f".//{W}body/{W}p"):
        text = "".join(t.text or "" for t in p.findall(f".//{W}t"))
        style = p.find(f"./{W}pPr/{W}pStyle")
        rows.append((text, "" if style is None else style.get(f"{W}val", "")))
    return rows


def _text(path: Path) -> list[str]:
    return [text for text, _style in _paragraphs(path) if text]


def _write_minimal_docx(
    path: Path,
    *,
    root_relationships: str | None = None,
    document_extra: str = "",
    document_relationships: str = "",
    extra_parts: dict[str, bytes] | None = None,
) -> None:
    document_xml = (
        f'<w:document xmlns:w="{W_NS}" xmlns:r="{R_NS}">'
        "<w:body><w:p><w:r><w:t>Body</w:t></w:r></w:p>"
        f"{document_extra}"
        "<w:sectPr /></w:body></w:document>"
    ).encode()
    parts = {
        "[Content_Types].xml": (
            b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            b'<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            b'<Default Extension="xml" ContentType="application/xml"/>'
            b'<Default Extension="png" ContentType="image/png"/>'
            b'<Override PartName="/word/document.xml" '
            b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            b"</Types>"
        ),
        "_rels/.rels": (
            f'<Relationships xmlns="{REL_NS}">'
            + (
                root_relationships
                if root_relationships is not None
                else f'<Relationship Id="rId1" Type="{OFFICE_DOCUMENT_REL}" Target="word/document.xml"/>'
            )
            + "</Relationships>"
        ).encode(),
        "word/document.xml": document_xml,
        "word/_rels/document.xml.rels": (
            f'<Relationships xmlns="{REL_NS}">{document_relationships}</Relationships>'
        ).encode(),
    }
    parts.update(extra_parts or {})
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, payload in parts.items():
            zf.writestr(name, payload)


def test_single_input_merge_preserves_document_structure_after_package_rewrite(tmp_path: Path) -> None:
    image = tmp_path / "one.png"
    _write_png(image, (255, 0, 0))
    source = tmp_path / "source.docx"
    output = tmp_path / "merged.docx"
    _write_docx(source, [("h1", "Chapter 1"), ("p", "Body text")], image=image)

    summary = merge_docx((source,), output)

    assert summary.output == output.resolve()
    assert _paragraphs(output) == _paragraphs(source)
    assert validate_docx_package(output).media_parts == validate_docx_package(source).media_parts


def test_merge_multiple_parts_preserves_text_media_and_relationship_integrity(tmp_path: Path) -> None:
    image1 = tmp_path / "one.png"
    image2 = tmp_path / "two.png"
    _write_png(image1, (255, 0, 0))
    _write_png(image2, (0, 0, 255))
    first = tmp_path / "part1.docx"
    second = tmp_path / "part2.docx"
    output = tmp_path / "merged.docx"
    _write_docx(first, [("h1", "Chapter 1"), ("p", "First body")], image=image1)
    _write_docx(second, [("h1", "Chapter 2"), ("p", "Second body")], image=image2)

    summary = merge_docx((first, second), output)

    assert _text(output) == ["Chapter 1", "First body", "Chapter 2", "Second body"]
    assert summary.validation.media_parts == 2
    assert summary.validation.relationships >= 2
    assert summary.validation.relationship_refs >= 2
    validate_docx_package(output)


def test_merge_with_outline_parts_inserts_h1_parts_and_demotes_chapters_to_h2(tmp_path: Path) -> None:
    first = tmp_path / "part1.docx"
    second = tmp_path / "part2.docx"
    output = tmp_path / "merged.docx"
    _write_docx(first, [("h1", "Chapter 1. One"), ("p", "First body")])
    _write_docx(second, [("h1", "Chapter 1. Two"), ("p", "Second body")])

    summary = merge_docx(
        (first, second),
        output,
        parts=(
            PartBoundary("Part 1", "Chapter 1. One"),
            PartBoundary("Part 2", "Chapter 1. Two"),
        ),
    )

    assert summary.outline is not None
    assert summary.outline.inserted_parts == 2
    assert summary.outline.demoted_headings == 2
    assert _paragraphs(output)[:6] == [
        ("Part 1", "Heading1"),
        ("Chapter 1. One", "Heading2"),
        ("First body", ""),
        ("Part 2", "Heading1"),
        ("Chapter 1. Two", "Heading2"),
        ("Second body", ""),
    ]


def test_merge_rejects_outline_parts_for_single_input(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, [("h1", "Chapter 1"), ("p", "Body")])

    with pytest.raises(DocxMergeError, match="multiple source DOCX"):
        merge_docx(
            (source,),
            tmp_path / "merged.docx",
            parts=(PartBoundary("Part 1", "Chapter 1"),),
        )


def test_merge_rejects_outline_part_count_mismatch(tmp_path: Path) -> None:
    first = tmp_path / "part1.docx"
    second = tmp_path / "part2.docx"
    _write_docx(first, [("h1", "Chapter 1")])
    _write_docx(second, [("h1", "Chapter 2")])

    with pytest.raises(DocxMergeError, match="once per input DOCX"):
        merge_docx(
            (first, second),
            tmp_path / "merged.docx",
            parts=(PartBoundary("Part 1", "Chapter 1"),),
        )


def test_docx_merge_cli_smoke_with_temp_docx_fixtures(tmp_path: Path) -> None:
    first = tmp_path / "part1.docx"
    second = tmp_path / "part2.docx"
    output = tmp_path / "merged.docx"
    _write_docx(first, [("h1", "Chapter 1"), ("p", "First")])
    _write_docx(second, [("h1", "Chapter 2"), ("p", "Second")])

    rc = cli.main([
        "docx",
        "merge",
        str(first),
        str(second),
        "--out",
        str(output),
    ])

    assert rc == 0
    assert _text(output) == ["Chapter 1", "First", "Chapter 2", "Second"]
    validate_docx_package(output)


def test_merge_three_part_source_fixture_preserves_media_and_outline(tmp_path: Path) -> None:
    images = [tmp_path / f"part{index}.png" for index in range(1, 4)]
    for index, image in enumerate(images):
        _write_png(image, (index * 80, 0, 255 - index * 80))
    first = tmp_path / "part1.docx"
    second = tmp_path / "part2.docx"
    third = tmp_path / "part3.docx"
    _write_docx(first, [("h1", "Chapter 1. One"), ("p", "First body")], images[0])
    _write_docx(second, [("h1", "Chapter 1. Two"), ("p", "Second body")], images[1])
    _write_docx(third, [("h1", "Chapter 1. Three"), ("p", "Third body")], images[2])
    output = tmp_path / "merged.docx"

    summary = merge_docx(
        (first, second, third),
        output,
        parts=(
            PartBoundary("Part 1", "Chapter 1. One"),
            PartBoundary("Part 2", "Chapter 1. Two"),
            PartBoundary("Part 3", "Chapter 1. Three"),
        ),
    )

    expected_text = [
        "Part 1",
        *_text(first),
        "Part 2",
        *_text(second),
        "Part 3",
        *_text(third),
    ]
    assert _text(output) == expected_text
    assert summary.validation.media_parts == 3
    assert validate_docx_package(output).media_parts == 3
    assert _paragraphs(output)[:2] == [
        ("Part 1", "Heading1"),
        ("Chapter 1. One", "Heading2"),
    ]


def test_package_validation_accepts_percent_encoded_relationship_target(tmp_path: Path) -> None:
    docx = tmp_path / "encoded.docx"
    _write_minimal_docx(
        docx,
        document_relationships=(
            f'<Relationship Id="rIdImage" Type="{R_NS}/image" Target="media/image%201.png"/>'
        ),
        extra_parts={"word/media/image 1.png": b"image"},
    )

    validate_docx_package(docx)


def test_package_validation_rejects_external_target_without_target_mode(tmp_path: Path) -> None:
    docx = tmp_path / "external-target.docx"
    _write_minimal_docx(
        docx,
        document_relationships=(
            f'<Relationship Id="rIdImage" Type="{R_NS}/image" '
            'Target="https://example.test/word/media/image1.png"/>'
        ),
        extra_parts={"word/media/image1.png": b"image"},
    )

    with pytest.raises(DocxMergeError, match="external without TargetMode=External"):
        validate_docx_package(docx)


def test_package_validation_rejects_invalid_target_mode(tmp_path: Path) -> None:
    docx = tmp_path / "invalid-target-mode.docx"
    _write_minimal_docx(
        docx,
        document_relationships=(
            f'<Relationship Id="rIdImage" Type="{R_NS}/image" '
            'Target="media/image1.png" TargetMode="external"/>'
        ),
        extra_parts={"word/media/image1.png": b"image"},
    )

    with pytest.raises(DocxMergeError, match="invalid TargetMode 'external'"):
        validate_docx_package(docx)


def test_package_validation_rejects_misplaced_nested_relationship_part(tmp_path: Path) -> None:
    docx = tmp_path / "misplaced-rels.docx"
    _write_minimal_docx(
        docx,
        extra_parts={
            "word/charts/chart1.xml": b"<chart />",
            "word/media/image1.png": b"image",
            "word/_rels/charts/chart1.xml.rels": (
                f'<Relationships xmlns="{REL_NS}">'
                f'<Relationship Id="rIdImage" Type="{R_NS}/image" Target="../media/image1.png"/>'
                "</Relationships>"
            ).encode(),
        },
    )

    with pytest.raises(DocxMergeError, match="unexpected relationships part path"):
        validate_docx_package(docx)


def test_package_validation_accepts_root_level_relationship_part(tmp_path: Path) -> None:
    docx = tmp_path / "root-level-rels.docx"
    _write_minimal_docx(
        docx,
        extra_parts={
            "custom.xml": f'<custom xmlns:r="{R_NS}" r:id="rIdCustom" />'.encode(),
            "custom-target.xml": b"<customTarget />",
            "_rels/custom.xml.rels": (
                f'<Relationships xmlns="{REL_NS}">'
                '<Relationship Id="rIdCustom" '
                f'Type="{R_NS}/hyperlink" Target="custom-target.xml"/>'
                "</Relationships>"
            ).encode(),
        },
    )

    validate_docx_package(docx)


def test_package_validation_rejects_missing_root_office_document_relationship(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "missing-root-office-doc.docx"
    _write_minimal_docx(docx, root_relationships="")

    with pytest.raises(DocxMergeError, match="root officeDocument relationship"):
        validate_docx_package(docx)


def test_package_validation_rejects_embed_pointing_to_external_relationship(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "external-embed.docx"
    _write_minimal_docx(
        docx,
        document_extra='<w:pict r:embed="rIdImage" />',
        document_relationships=(
            f'<Relationship Id="rIdImage" Type="{R_NS}/image" '
            'Target="https://example.test/image.png" TargetMode="External"/>'
        ),
    )

    with pytest.raises(DocxMergeError, match="r:embed=rIdImage pointing to an external relationship"):
        validate_docx_package(docx)


def test_package_validation_rejects_embed_pointing_to_non_embeddable_relationship(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "non-embeddable-embed.docx"
    _write_minimal_docx(
        docx,
        document_extra='<w:pict r:embed="rIdLink" />',
        document_relationships=(
            f'<Relationship Id="rIdLink" Type="{R_NS}/hyperlink" Target="media/image1.png"/>'
        ),
        extra_parts={"word/media/image1.png": b"image"},
    )

    with pytest.raises(DocxMergeError, match="non-embeddable relationship type"):
        validate_docx_package(docx)


def test_package_validation_counts_each_xml_relationship_attribute(tmp_path: Path) -> None:
    docx = tmp_path / "relationship-refs.docx"
    _write_minimal_docx(
        docx,
        document_extra=(
            '<w:pict r:embed="rIdImage" />'
            '<w:pict r:embed="rIdImage" />'
        ),
        document_relationships=(
            f'<Relationship Id="rIdImage" Type="{R_NS}/image" Target="media/image1.png"/>'
        ),
        extra_parts={"word/media/image1.png": b"image"},
    )

    assert validate_docx_package(docx).relationship_refs == 2


def test_merge_rejects_empty_inputs(tmp_path: Path) -> None:
    with pytest.raises(DocxMergeError, match="at least one input"):
        merge_docx((), tmp_path / "out.docx")


def test_merge_rejects_missing_input(tmp_path: Path) -> None:
    with pytest.raises(DocxMergeError, match="DOCX not found"):
        merge_docx((tmp_path / "missing.docx",), tmp_path / "out.docx")


def test_merge_rejects_duplicate_inputs(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, [("p", "Body")])

    with pytest.raises(DocxMergeError, match="duplicate input DOCX"):
        merge_docx((source, source), tmp_path / "out.docx")


def test_merge_rejects_invalid_output_parent(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, [("p", "Body")])

    with pytest.raises(DocxMergeError, match="output parent does not exist"):
        merge_docx((source,), tmp_path / "missing" / "out.docx")
